import os
import json
import boto3
import stripe
import requests
from datetime import datetime
from werkzeug.utils import secure_filename
import uuid
import io
from PyPDF2 import PdfReader
from docx import Document
import threading

# Load environment variables from .env file
# Always use manual loading to ensure environment variables are set in subprocess context
env_file = '/var/www/lexi/.env'
if os.path.exists(env_file):
    with open(env_file, 'r') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, val = line.split('=', 1)
                key = key.strip()
                val = val.strip().strip('"\'')
                # Only set if not already in environment (prefer shell-set vars)
                if key not in os.environ:
                    os.environ[key] = val

# Also try python-dotenv as backup
try:
    from dotenv import load_dotenv
    load_dotenv('/var/www/lexi/.env')
except ImportError:
    pass

# Productie Stripe key heeft voorrang over test key
stripe.api_key = os.getenv('STRIPE_SECRET_KEY_PROD') or os.getenv('STRIPE_SECRET_KEY', '')

class MemgraphDeepSeekService:
    """
    Singleton Memgraph + DeepSeek V3 RAG Service

    Open Source Stack:
    - Knowledge Graph: Memgraph (for relationship-aware retrieval)
    - LLM: DeepSeek V3 (cost-effective, high-quality)
    - Embeddings: sentence-transformers (self-hosted)

    CRITICAL for production deployment:
    - Uses thread-safe singleton pattern
    - Graph-based context retrieval (better than vector-only)
    - 37x cheaper than GPT-4 for comparable quality
    """
    _instance = None
    _lock = threading.Lock()

    def __new__(cls):
        """Thread-safe singleton implementation"""
        if cls._instance is None:
            with cls._lock:
                # Double-check locking pattern
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        """Initialize only once per process (singleton pattern)"""
        # Skip if already initialized
        if self._initialized:
            return

        self.memgraph = None
        self.embedding_model = None
        self.voyage_client = None
        self.enabled = False  # Always initialize to prevent AttributeError

        try:
            # Import dependencies
            try:
                from gqlalchemy import Memgraph
                from sentence_transformers import SentenceTransformer
                from voyageai import Client as VoyageClient
            except ImportError as e:
                print(f"❌ Missing dependencies: {e}")
                print("Install with: pip install gqlalchemy sentence-transformers torch")
                print("⚠️  Memgraph+DeepSeek RAG will be DISABLED until dependencies are installed")
                # Don't raise - allow app to start without RAG
                self._initialized = True
                return

            # Memgraph connection
            memgraph_host = os.getenv('MEMGRAPH_HOST', 'localhost')
            memgraph_port = int(os.getenv('MEMGRAPH_PORT', 7687))
            memgraph_username = os.getenv('MEMGRAPH_USERNAME', None)
            memgraph_password = os.getenv('MEMGRAPH_PASSWORD', None)

            # Build connection kwargs - only include username/password if explicitly set
            mg_kwargs = {'host': memgraph_host, 'port': memgraph_port}
            if memgraph_username:
                mg_kwargs['username'] = memgraph_username
            if memgraph_password:
                mg_kwargs['password'] = memgraph_password

            self.memgraph = Memgraph(**mg_kwargs)

            # Test connection
            try:
                list(self.memgraph.execute_and_fetch("RETURN 1 as test"))
                print(f"✓ Memgraph connected: {memgraph_host}:{memgraph_port}")
            except Exception as e:
                print(f"⚠️  Memgraph not available: {e}")
                print(f"   Start Memgraph: docker run -p 7687:7687 memgraph/memgraph-platform")
                raise

            # Load embedding model (multilingual, optimized for Dutch - LEGACY FALLBACK)
            embedding_model_name = os.getenv('EMBEDDING_MODEL', 'intfloat/multilingual-e5-large')
            print(f"Loading embedding model: {embedding_model_name}...")
            self.embedding_model = SentenceTransformer(embedding_model_name)
            print(f"✓ Embedding model loaded (dim={self.embedding_model.get_sentence_embedding_dimension()})")

            # Voyage AI (Production Grade Legal Embeddings)
            voyage_api_key = os.getenv('VOYAGE_AI_API_KEY') or os.getenv('VOYAGE_API_KEY')
            if voyage_api_key:
                try:
                    self.voyage_client = VoyageClient(api_key=voyage_api_key)
                    self.voyage_model = os.getenv('VOYAGE_AI_MODEL', 'voyage-law-2')
                    print(f"✓ Voyage AI initialized (model: {self.voyage_model})")
                except Exception as e:
                    print(f"⚠️  Voyage AI init failed: {e} - using fallback embeddings")
                    self.voyage_client = None
            else:
                print("⚠️  VOYAGE_AI_API_KEY not set - using fallback embeddings")

            # DeepSeek API configuration
            self.deepseek_api_key = os.getenv('DEEPSEEK_API_KEY')
            self.deepseek_api_url = os.getenv('DEEPSEEK_API_URL', 'https://api.deepseek.com/v1/chat/completions')
            self.deepseek_model = os.getenv('DEEPSEEK_MODEL', 'deepseek-chat')

            if not self.deepseek_api_key:
                print("⚠️  DEEPSEEK_API_KEY not set - RAG will not work")
                print("   Get API key: https://platform.deepseek.com/")
                raise ValueError("DEEPSEEK_API_KEY required")

            # Generieke fallback system instruction (wordt overschreven door cao_config.py)
            self.system_instruction = """Je bent Lexi - Expert Loonadministrateur voor uitzendbureaus.

KERN INSTRUCTIES:
- Gebruik de kennisbank (CAO documenten) om de beste antwoorden te geven
- Geef concrete, bruikbare adviezen
- Wees transparant over bronnen
- Gebruik ALLEEN de context die je krijgt - geen externe bronnen

ANTWOORD STRUCTUUR:
1. BESLUIT: [Duidelijke conclusie]
2. BASIS: [Gevonden in CAO documenten + artikel referenties]
3. ACTIE: [Concrete stappen]

DOCUMENT GENERATIE (ARTIFACTS):
Wanneer je gevraagd wordt om een document te maken (contract, brief, formulier, etc.), gebruik dit format:

```artifact:document title:Naam van het document
[Volledige inhoud van het document hier]
```

Voorbeeld:
```artifact:contract title:Arbeidsovereenkomst Uitzendkracht
ARBEIDSOVEREENKOMST

Tussen: [Werkgever]
En: [Werknemer]

Artikel 1 - Functie
...
```

Dit creëert automatisch een downloadbaar document voor de gebruiker.

Gebruik alle beschikbare documenten optimaal. Je bent expert-niveau - vertrouw op je analyse."""

            self.enabled = True
            print(f"✓ Memgraph + DeepSeek RAG initialized (Voyage AI: {'✓' if self.voyage_client else '⚠️  fallback'})")
            self._initialized = True
            print("✓ Memgraph + DeepSeek RAG initialized successfully (singleton)")

        except Exception as e:
            print(f"❌ Memgraph+DeepSeek initialization failed: {e}")
            import traceback
            traceback.print_exc()
            self.enabled = False
            self._initialized = True

    def chat(self, message, conversation_history=None, system_instruction=None):
        """
        Chat with RAG context from Memgraph + DeepSeek V3

        Args:
            message: User's question
            conversation_history: List of previous messages
            system_instruction: Custom system instruction (overrides default)

        Returns:
            str: AI response
        """
        if not self.enabled:
            return "Lexi is momenteel niet beschikbaar. Controleer de Memgraph en DeepSeek configuratie."

        try:
            # 1. Generate embedding for user query (Voyage AI preferred)
            if self.voyage_client:
                try:
                    result = self.voyage_client.embed([message], model=self.voyage_model)
                    embedding = result.embeddings[0]
                except Exception as e:
                    print(f"⚠️  Voyage AI embedding failed: {e}, using fallback")
                    embedding = self.embedding_model.encode(message, convert_to_tensor=False).tolist()
            else:
                embedding = self.embedding_model.encode(message, convert_to_tensor=False).tolist()

            # 2. Query Memgraph for relevant context (Voyage AI semantic matching)
            # Extract keywords from query for better matching
            query_keywords = message.lower().split()

            try:
                # Try vector search first (if available)
                cypher = """
                CALL vector.search('cao_embeddings', $embedding, 10)
                YIELD node, similarity
                WHERE similarity > 0.65
                WITH node, similarity
                ORDER BY similarity DESC
                LIMIT 5

                OPTIONAL MATCH (node)-[:REFERENCES|APPLIES_TO|EXCEPTION]-(related:Artikel)
                WHERE related.cao = node.cao OR related.cao IS NULL

                RETURN
                    node.text AS text,
                    node.article_number AS article,
                    node.cao AS cao,
                    similarity,
                    collect(DISTINCT related.title)[0..2] AS related_articles
                """

                results = list(self.memgraph.execute_and_fetch(
                    cypher,
                    {'embedding': embedding}
                ))
            except Exception as e:
                # Fallback: Use semantic title/section matching from Artikel nodes
                print(f"⚠️  Vector search unavailable: {e}, using title/section matching")
                results = []

                try:
                    # Simple fallback: search by title and section keywords
                    all_articles = list(self.memgraph.execute_and_fetch(
                        "MATCH (cao:CAO)-[:CONTAINS_ARTIKEL]->(a:Artikel) RETURN a.title AS text, a.number AS article, cao.name AS cao, a.section AS section, a.tags AS tags"
                    ))

                    # Score articles based on keyword matching
                    scored = []
                    for art in all_articles:
                        score = 0
                        title_lower = (art['text'] or '').lower()
                        section_lower = (art['section'] or '').lower()
                        tags_lower = (art['tags'] or '').lower()

                        for kw in query_keywords:
                            if len(kw) > 2:  # Skip small words
                                if kw in title_lower:
                                    score += 3
                                if kw in section_lower:
                                    score += 2
                                if kw in tags_lower:
                                    score += 2

                        if score > 0:
                            art['similarity'] = score
                            art['related_articles'] = [art['section']]
                            scored.append(art)

                    # Sort by score and return top 5
                    results = sorted(scored, key=lambda x: x['similarity'], reverse=True)[:5]

                except Exception as e2:
                    print(f"⚠️  Keyword matching failed: {e2}")
                    results = []

            # 3. Build context from results
            context_parts = []
            if results:
                for r in results:
                    article_ref = f"[{r['cao']} Artikel {r['article']}]" if r.get('article') else ""
                    context_parts.append(f"{article_ref} {r['text']}")

                    # Add related articles for richer context
                    if r.get('related_articles'):
                        for related in r['related_articles']:
                            if related and len(related) > 50:
                                context_parts.append(f"  → Gerelateerd: {related[:300]}...")

            context = "\n\n".join(context_parts) if context_parts else "Geen relevante documenten gevonden."

            # 4. Build messages for DeepSeek
            instruction_text = system_instruction if system_instruction else self.system_instruction

            messages = [
                {
                    "role": "system",
                    "content": f"{instruction_text}\n\n---\n\nRELEVANTE CAO CONTEXT:\n{context}"
                }
            ]

            # Add conversation history
            if conversation_history:
                for msg in conversation_history:
                    messages.append({
                        "role": msg.get('role', 'user'),
                        "content": msg.get('content', '')
                    })

            # Add current message
            messages.append({
                "role": "user",
                "content": message
            })

            # 5. Call DeepSeek API with streaming
            import httpx

            full_response = ""

            with httpx.stream(
                'POST',
                self.deepseek_api_url,
                headers={
                    'Authorization': f'Bearer {self.deepseek_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    'model': self.deepseek_model,
                    'messages': messages,
                    'stream': True,
                    'temperature': 1.0,
                    'max_tokens': 4096
                },
                timeout=120
            ) as response:
                response.raise_for_status()

                for line in response.iter_lines():
                    if not line:
                        continue

                    if line.startswith('data: '):
                        data_str = line[6:]  # Remove 'data: ' prefix

                        if data_str.strip() == '[DONE]':
                            break

                        try:
                            chunk = json.loads(data_str)
                            if chunk.get('choices') and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', {})
                                content = delta.get('content')
                                if content:
                                    full_response += content
                        except json.JSONDecodeError:
                            continue

            return full_response if full_response else "Geen response ontvangen van AI."

        except httpx.HTTPStatusError as e:
            print(f"❌ DeepSeek API error: {e.response.status_code} - {e.response.text}")
            return f"Er ging iets mis bij het verwerken van je vraag (API error: {e.response.status_code})."

        except Exception as e:
            print(f"❌ Chat error: {e}")
            import traceback
            traceback.print_exc()
            return "Er ging iets mis bij het verwerken van je vraag. Probeer het opnieuw."


class DeepSeekR1Client:
    """
    DeepSeek R1 (Reasoning Model) Client for Automatic CAO Structure Analysis

    Specialized for analyzing document chunks and extracting:
    - CAO metadata (name, version, type, sector, date)
    - Article identification and proper numbering
    - Article relationships and cross-references
    - Document structure hierarchy

    Uses DeepSeek R1 (deepseek-reasoner) for structured, reasoning-based analysis.
    This model thinks deeply before answering, ensuring accurate document indexing.

    Critical for GraphRAG: Proper extraction ensures high-quality graph structure
    and better retrieval relevance.
    """
    _instance = None
    _lock = threading.Lock()

    def __new__(cls):
        """Thread-safe singleton implementation"""
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        """Initialize DeepSeek R1 client"""
        if self._initialized:
            return

        self.deepseek_api_key = os.getenv('DEEPSEEK_API_KEY')
        self.deepseek_api_url = os.getenv('DEEPSEEK_API_URL', 'https://api.deepseek.com/v1/chat/completions')
        self.deepseek_r1_model = 'deepseek-reasoner'  # Reasoning model

        if not self.deepseek_api_key:
            print("⚠️  DEEPSEEK_API_KEY not set - R1 analysis will not work")
            self.enabled = False
        else:
            self.enabled = True
            print(f"✓ DeepSeek R1 Client initialized (model: {self.deepseek_r1_model})")

        self._initialized = True

    def analyze_cao_structure(self, chunks: list, document_name: str, cao_type: str = None) -> dict:
        """
        Analyze document chunks using DeepSeek R1 to extract CAO structure

        Args:
            chunks: List of text chunks from document
            document_name: Name of the uploaded document
            cao_type: Type of CAO (NBBU, ABU, Glastuinbouw, etc.) if known

        Returns:
            {
                'success': bool,
                'cao_metadata': {
                    'name': str,
                    'type': str,
                    'version': str,
                    'effective_date': str,
                    'sector': str,
                    'description': str
                },
                'artikelen': [
                    {
                        'article_number': str,
                        'title': str,
                        'content_preview': str,
                        'chunk_indices': [int],
                        'section': str,
                        'tags': [str]
                    }
                ],
                'relaties': [
                    {
                        'source_article': str,
                        'target_article': str,
                        'relation_type': str,
                        'description': str
                    }
                ],
                'validation': {
                    'total_articles_found': int,
                    'coverage_percentage': float,
                    'warnings': [str]
                },
                'tokens_used': int,
                'processing_time': float
            }
        """
        if not self.enabled:
            return {
                'success': False,
                'error': 'DeepSeek R1 client not initialized',
                'tokens_used': 0
            }

        import time
        start_time = time.time()

        try:
            # Prepare chunks summary for analysis
            chunks_text = "\n\n---CHUNK SEPARATOR---\n\n".join(
                f"[CHUNK {i}]\n{chunk[:1000]}"
                for i, chunk in enumerate(chunks[:50])  # Limit to first 50 chunks for analysis
            )

            # Create analysis prompt for R1 reasoning
            analysis_prompt = f"""Analyze this legal document (CAO/collective labor agreement) and extract its structure.

DOCUMENT: {document_name}
CAO TYPE: {cao_type or 'Unknown'}

CONTENT CHUNKS:
{chunks_text}

ANALYZE AND EXTRACT:

1. **CAO METADATA**:
   - Official name
   - Type (NBBU, ABU, Glastuinbouw, etc.)
   - Version number if present
   - Effective date range
   - Sector/Industry
   - Brief description

2. **ARTICLES** (Extract all article numbers and info):
   For each article found, provide:
   - Article number (e.g., "Article 1", "Artikel 1.2", etc.)
   - Article title
   - Which chunks contain this article (by chunk number)
   - Section it belongs to
   - Tags (e.g., "wages", "leave", "termination", etc.)

3. **RELATIONSHIPS**:
   - Which articles reference or depend on each other
   - Cross-references and related topics
   - Hierarchical structure (if any)

4. **VALIDATION**:
   - Estimate total articles in document (including parts you haven't seen)
   - Estimated coverage percentage from chunks provided
   - Any warnings or issues with document structure

RESPOND IN STRICT JSON FORMAT:
{{
    "cao_metadata": {{
        "name": "...",
        "type": "...",
        "version": "...",
        "effective_date": "...",
        "sector": "...",
        "description": "..."
    }},
    "artikelen": [
        {{"article_number": "...", "title": "...", "chunk_indices": [...], "section": "...", "tags": [...]}}
    ],
    "relaties": [
        {{"source_article": "...", "target_article": "...", "relation_type": "...", "description": "..."}}
    ],
    "validation": {{
        "total_articles_estimated": 123,
        "coverage_percentage": 45.5,
        "warnings": ["..."]
    }}
}}"""

            # Call DeepSeek R1 API
            import httpx

            response_text = ""
            total_tokens = 0

            with httpx.stream(
                'POST',
                self.deepseek_api_url,
                headers={
                    'Authorization': f'Bearer {self.deepseek_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    'model': self.deepseek_r1_model,
                    'messages': [
                        {
                            'role': 'user',
                            'content': analysis_prompt
                        }
                    ],
                    'temperature': 0.1,  # Low temperature for consistent reasoning
                    'max_tokens': 8000,  # Allow detailed analysis
                    'stream': True
                },
                timeout=300  # 5 minute timeout for reasoning
            ) as response:
                response.raise_for_status()

                for line in response.iter_lines():
                    if not line:
                        continue

                    if line.startswith('data: '):
                        data_str = line[6:]

                        if data_str.strip() == '[DONE]':
                            break

                        try:
                            chunk = json.loads(data_str)
                            if chunk.get('choices') and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', {})
                                content = delta.get('content')
                                if content:
                                    response_text += content

                                # Track token usage
                                if chunk.get('usage'):
                                    total_tokens = chunk['usage'].get('total_tokens', 0)
                        except json.JSONDecodeError:
                            continue

            # Parse R1 response
            # Extract JSON from response (R1 may include reasoning before JSON)
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)

            if not json_match:
                return {
                    'success': False,
                    'error': 'Could not parse R1 response as JSON',
                    'raw_response': response_text[:500],
                    'tokens_used': total_tokens,
                    'processing_time': time.time() - start_time
                }

            analysis_result = json.loads(json_match.group())

            # Validate and enrich result
            result = {
                'success': True,
                'cao_metadata': analysis_result.get('cao_metadata', {}),
                'artikelen': analysis_result.get('artikelen', []),
                'relaties': analysis_result.get('relaties', []),
                'validation': analysis_result.get('validation', {
                    'total_articles_estimated': len(analysis_result.get('artikelen', [])),
                    'coverage_percentage': 100.0,
                    'warnings': []
                }),
                'tokens_used': total_tokens,
                'processing_time': time.time() - start_time
            }

            print(f"✓ R1 Analysis complete: {len(result['artikelen'])} articles extracted, {total_tokens} tokens used")
            return result

        except httpx.HTTPStatusError as e:
            error_msg = f"DeepSeek API error: {e.response.status_code}"
            print(f"❌ {error_msg}")
            try:
                error_detail = e.response.json()
                print(f"   Details: {error_detail}")
            except:
                print(f"   Body: {e.response.text[:200]}")

            return {
                'success': False,
                'error': error_msg,
                'tokens_used': 0,
                'processing_time': time.time() - start_time
            }

        except Exception as e:
            print(f"❌ R1 Analysis error: {e}")
            import traceback
            traceback.print_exc()

            return {
                'success': False,
                'error': str(e),
                'tokens_used': 0,
                'processing_time': time.time() - start_time
            }

    def validate_analysis(self, analysis: dict) -> tuple[bool, list]:
        """
        Validate R1 analysis result

        Args:
            analysis: Result from analyze_cao_structure()

        Returns:
            (is_valid, warnings)
        """
        warnings = []

        if not analysis.get('success'):
            return False, [f"Analysis failed: {analysis.get('error', 'Unknown error')}"]

        # Check CAO metadata
        cao_meta = analysis.get('cao_metadata', {})
        if not cao_meta.get('name'):
            warnings.append("CAO name not found")

        # Check articles
        artikelen = analysis.get('artikelen', [])
        if not artikelen:
            warnings.append("No articles extracted")
        else:
            # Validate article numbers
            for artikel in artikelen:
                if not artikel.get('article_number'):
                    warnings.append(f"Article missing number: {artikel}")
                if not artikel.get('title'):
                    warnings.append(f"Article {artikel.get('article_number', '?')} missing title")

        # Validation metrics
        validation = analysis.get('validation', {})
        coverage = validation.get('coverage_percentage', 0)
        if coverage < 30:
            warnings.append(f"Low coverage: only {coverage}% of document analyzed")

        is_valid = len(warnings) < 3  # Allow up to 2 warnings
        return is_valid, warnings


# Global R1 client instance
_r1_client_instance = None

def get_r1_client() -> DeepSeekR1Client:
    """Get or create global DeepSeek R1 client"""
    global _r1_client_instance
    if _r1_client_instance is None:
        _r1_client_instance = DeepSeekR1Client()
    return _r1_client_instance


class S3Service:
    """
    Singleton S3 Service
    
    CRITICAL for production deployment:
    - Uses thread-safe singleton pattern to prevent client recreation
    - Prevents boto3 client creation crashes in gunicorn workers
    - Ensures single S3 client instance across all requests
    """
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        """Thread-safe singleton implementation"""
        if cls._instance is None:
            with cls._lock:
                # Double-check locking pattern
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        """Initialize only once per process (singleton pattern)"""
        # Skip if already initialized
        if self._initialized:
            return
            
        self.endpoint = os.getenv('S3_ENDPOINT_URL')
        self.bucket = os.getenv('S3_BUCKET_NAME')
        self.access_key = os.getenv('S3_ACCESS_KEY')
        self.secret_key = os.getenv('S3_SECRET_KEY')
        
        if all([self.endpoint, self.bucket, self.access_key, self.secret_key]):
            try:
                self.s3_client = boto3.client(
                    's3',
                    endpoint_url=self.endpoint,
                    aws_access_key_id=self.access_key,
                    aws_secret_access_key=self.secret_key
                )
                self.enabled = True
                self._initialized = True
                print("✓ S3 Service initialized successfully (singleton)")
            except Exception as e:
                print(f"S3 initialization failed: {e}")
                self.enabled = False
                self._initialized = True
        else:
            self.enabled = False
            self._initialized = True
            print("S3 Service disabled (missing credentials)")
    
    def upload_file(self, file, tenant_id, folder='uploads'):
        if not self.enabled:
            return None
        
        try:
            filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4()}_{filename}"
            s3_key = f"{folder}/tenant_{tenant_id}/{unique_filename}"
            
            self.s3_client.upload_fileobj(
                file,
                self.bucket,
                s3_key,
                ExtraArgs={'ContentType': file.content_type or 'application/octet-stream'}
            )
            
            return s3_key
        except Exception as e:
            print(f"S3 upload error: {e}")
            return None
    
    def upload_content(self, content, filename, tenant_id, folder='artifacts'):
        if not self.enabled:
            return None
        
        try:
            unique_filename = f"{uuid.uuid4()}_{filename}"
            s3_key = f"{folder}/tenant_{tenant_id}/{unique_filename}"
            
            self.s3_client.put_object(
                Bucket=self.bucket,
                Key=s3_key,
                Body=content.encode('utf-8') if isinstance(content, str) else content,
                ContentType='text/plain'
            )
            
            return s3_key
        except Exception as e:
            print(f"S3 upload content error: {e}")
            return None
    
    def download_file_content(self, s3_key, mime_type=None):
        if not self.enabled:
            return None, "S3 niet geconfigureerd"
        
        try:
            response = self.s3_client.get_object(Bucket=self.bucket, Key=s3_key)
            content_bytes = response['Body'].read()
            
            if mime_type == 'application/pdf':
                try:
                    pdf_file = io.BytesIO(content_bytes)
                    pdf_reader = PdfReader(pdf_file)
                    text_content = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text)
                    
                    if not text_content:
                        return None, "PDF bevat geen leesbare tekst (mogelijk scan of beveiligd)"
                    
                    return '\n'.join(text_content), None
                except Exception as e:
                    return None, f"Kon PDF niet lezen: {str(e)}"
            
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    docx_file = io.BytesIO(content_bytes)
                    doc = Document(docx_file)
                    text_content = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_content.append(' | '.join(row_text))
                    
                    if not text_content:
                        return None, "DOCX bevat geen leesbare tekst"
                    
                    return '\n'.join(text_content), None
                except Exception as e:
                    return None, f"Kon DOCX niet lezen: {str(e)}"
            
            if mime_type and 'text' in mime_type:
                try:
                    return content_bytes.decode('utf-8'), None
                except UnicodeDecodeError:
                    return content_bytes.decode('latin-1'), None
            
            try:
                return content_bytes.decode('utf-8'), None
            except UnicodeDecodeError:
                return None, "Kon bestand niet lezen. Upload alleen tekst, PDF of DOCX bestanden."
                
        except Exception as e:
            print(f"S3 download error: {e}")
            return None, f"Fout bij downloaden: {str(e)}"
    
    def get_file_url(self, s3_key, expiration=3600):
        if not self.enabled:
            return None
        
        try:
            url = self.s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': self.bucket, 'Key': s3_key},
                ExpiresIn=expiration
            )
            return url
        except Exception as e:
            print(f"S3 get URL error: {e}")
            return None
    
    def delete_file(self, s3_key):
        if not self.enabled:
            return False
        
        try:
            self.s3_client.delete_object(Bucket=self.bucket, Key=s3_key)
            return True
        except Exception as e:
            print(f"S3 delete error: {e}")
            return False
    
    def save_chat_messages(self, chat_id, tenant_id, messages):
        """Save chat messages to S3 as JSON"""
        if not self.enabled:
            return None
        
        try:
            s3_key = f"chats/tenant_{tenant_id}/chat_{chat_id}_messages.json"
            
            messages_data = json.dumps(messages, ensure_ascii=False, indent=2)
            
            self.s3_client.put_object(
                Bucket=self.bucket,
                Key=s3_key,
                Body=messages_data.encode('utf-8'),
                ContentType='application/json'
            )
            
            return s3_key
        except Exception as e:
            print(f"S3 save chat messages error: {e}")
            return None
    
    def get_chat_messages(self, s3_key):
        """Get chat messages from S3"""
        if not self.enabled:
            return []
        
        try:
            response = self.s3_client.get_object(Bucket=self.bucket, Key=s3_key)
            content = response['Body'].read().decode('utf-8')
            return json.loads(content)
        except self.s3_client.exceptions.NoSuchKey:
            return []
        except Exception as e:
            print(f"S3 get chat messages error: {e}")
            return []
    
    def get_messages(self, s3_key):
        """Get chat messages from S3 wrapped in messages dict"""
        if not self.enabled:
            return {'messages': []}
        
        try:
            response = self.s3_client.get_object(Bucket=self.bucket, Key=s3_key)
            content = response['Body'].read().decode('utf-8')
            messages = json.loads(content)
            return {'messages': messages}
        except self.s3_client.exceptions.NoSuchKey:
            return {'messages': []}
        except Exception as e:
            print(f"S3 get messages error: {e}")
            return {'messages': []}
    
    def append_chat_message(self, s3_key, chat_id, tenant_id, message):
        """Append a new message to existing chat in S3"""
        if not self.enabled:
            return False
        
        try:
            messages = self.get_chat_messages(s3_key) if s3_key else []
            messages.append(message)
            
            new_s3_key = self.save_chat_messages(chat_id, tenant_id, messages)
            return new_s3_key
        except Exception as e:
            print(f"S3 append chat message error: {e}")
            return None

class StripeService:
    @staticmethod
    def create_checkout_session(tenant_id, plan, success_url, cancel_url):
        try:
            price_id = os.getenv(f'STRIPE_PRICE_{plan.upper()}')
            if not price_id:
                if plan == 'starter':
                    price_id = 'price_starter'
                elif plan == 'professional':
                    price_id = 'price_professional'
                else:
                    price_id = 'price_enterprise'
            
            session = stripe.checkout.Session.create(
                payment_method_types=['card', 'ideal'],
                line_items=[{
                    'price': price_id,
                    'quantity': 1,
                }],
                mode='subscription',
                success_url=success_url,
                cancel_url=cancel_url,
                metadata={'tenant_id': tenant_id, 'plan': plan}
            )
            return session
        except Exception as e:
            print(f"Stripe checkout error: {e}")
            return None
    
    @staticmethod
    def create_customer_portal_session(customer_id, return_url):
        try:
            session = stripe.billing_portal.Session.create(
                customer=customer_id,
                return_url=return_url,
            )
            return session
        except Exception as e:
            print(f"Stripe portal error: {e}")
            return None

class EmailService:
    """
    Singleton Email Service
    
    CRITICAL for production deployment:
    - Uses thread-safe singleton pattern for consistency
    - Prevents multiple MailerSend API initializations
    - Ensures single email service instance across all requests
    """
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        """Thread-safe singleton implementation"""
        if cls._instance is None:
            with cls._lock:
                # Double-check locking pattern
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        """Initialize only once per process (singleton pattern)"""
        # Skip if already initialized
        if self._initialized:
            return
            
        self.api_key = os.getenv('MAILERSEND_API_KEY')
        self.from_email = os.getenv('FROM_EMAIL', 'noreply@trial-3vz9dle4n8z4kj50.mlsender.net')
        self.from_name = os.getenv('FROM_NAME', 'Lexi CAO Meester')
        self.enabled = bool(self.api_key)
        self.api_url = "https://api.mailersend.com/v1/email"
        
        # TEST_EMAIL_OVERRIDE: Route all emails to this address for layout testing
        self.test_email_override = os.getenv('TEST_EMAIL_OVERRIDE', '')
        
        self._initialized = True
        
        if self.enabled:
            print(f"✓ MailerSend HTTP API initialized: {self.from_name} <{self.from_email}> (singleton)")
            if self.test_email_override:
                print(f"⚠️  TEST MODE: All emails redirected to {self.test_email_override}")
    
    def send_email(self, to_email, subject, html_content):
        """Send email via MailerSend HTTP API (stable, production-ready)"""
        if not self.enabled:
            print(f"Email not sent (MailerSend not configured): {subject} to {to_email}")
            return False
        
        # Override recipient for testing if TEST_EMAIL_OVERRIDE is set
        original_to_email = to_email
        if self.test_email_override:
            to_email = self.test_email_override
            print(f"📧 TEST MODE: Redirecting email from {original_to_email} to {to_email}")
        
        try:
            # Strip HTML tags for plain text version
            import re
            text_content = re.sub('<[^<]+?>', '', html_content)
            
            # Build email payload for HTTP API
            payload = {
                "from": {
                    "email": self.from_email,
                    "name": self.from_name
                },
                "to": [
                    {
                        "email": to_email
                    }
                ],
                "subject": subject,
                "text": text_content,
                "html": html_content
            }
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
                "X-Requested-With": "XMLHttpRequest"
            }
            
            # Send email via HTTP POST
            response = requests.post(
                self.api_url,
                headers=headers,
                json=payload,
                timeout=10
            )
            
            if response.status_code == 202:
                print(f"✓ Email sent successfully to {to_email} (subject: {subject})")
                return True
            else:
                print(f"MailerSend error: Status {response.status_code}, Response: {response.text}")
                return False
                
        except Exception as e:
            print(f"MailerSend error: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def send_welcome_email(self, user, tenant, login_url):
        """Send welcome email after successful signup (branded template)"""
        subject = "Welkom bij Lexi CAO Meester!"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Welkom bij Lexi! 👋</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {user.first_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Je account is aangemaakt voor <strong style="color: #1a2332;">{tenant.company_name}</strong>. Lexi staat klaar om al je CAO vragen te beantwoorden!
                                    </p>
                                    <div style="background-color: #d4af37; border-radius: 8px; padding: 20px; margin: 24px 0; text-align: center;">
                                        <p style="margin: 0 0 12px 0; color: #1a2332; font-size: 18px; font-weight: 600;">Start nu met Lexi!</p>
                                        <a href="{login_url}" style="background: #1a2332; color: #d4af37; padding: 12px 32px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: 600;">
                                            Inloggen →
                                        </a>
                                    </div>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <h3 style="margin: 0 0 12px 0; color: #1a2332; font-size: 16px; font-weight: 600;">Wat kun je met Lexi?</h3>
                                        <ul style="margin: 0; padding-left: 20px; color: #374151; line-height: 1.8;">
                                            <li>CAO vragen stellen en directe antwoorden krijgen</li>
                                            <li>Documenten genereren (contracten, brieven)</li>
                                            <li>Bestanden uploaden voor analyse</li>
                                            <li>Chat geschiedenis doorzoeken</li>
                                        </ul>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong> - Jouw Expert CAO Assistent
                                    </p>
                                    <p style="margin: 0; color: #9ca3af; font-size: 12px;">
                                        Vragen? support@lexiai.nl
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(user.email, subject, html_content)
    
    def send_user_invitation_email(self, user, tenant, activation_url, admin_name):
        """Send invitation email with secure activation link (NO PASSWORD IN EMAIL!)"""
        subject = f"Je bent uitgenodigd voor Lexi CAO Meester bij {tenant.company_name}"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Welkom bij Lexi! 👋</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {user.first_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        {admin_name} heeft een account voor je aangemaakt bij <strong style="color: #1a2332;">{tenant.company_name}</strong>. 
                                        Je hebt nu toegang tot Lexi, jouw AI-assistent voor CAO-vragen.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <h3 style="margin: 0 0 12px 0; color: #1a2332; font-size: 18px; font-weight: 600;">🔐 Account Activeren</h3>
                                        <p style="margin: 0 0 16px 0; color: #374151; font-size: 14px; line-height: 1.6;">
                                            Klik op de knop hieronder om je account te activeren en je wachtwoord in te stellen.
                                        </p>
                                        <p style="margin: 0; color: #6b7280; font-size: 14px; line-height: 1.6;">
                                            <strong>Email:</strong> {user.email}
                                        </p>
                                    </div>
                                    <div style="background-color: #d4af37; border-radius: 8px; padding: 20px; margin: 24px 0; text-align: center;">
                                        <p style="margin: 0 0 12px 0; color: #1a2332; font-size: 18px; font-weight: 600;">Activeer je Account</p>
                                        <a href="{activation_url}" style="background: #1a2332; color: #d4af37; padding: 14px 40px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: 600; font-size: 16px;">
                                            Account Activeren →
                                        </a>
                                    </div>
                                    <div style="margin: 32px 0; padding: 24px; background-color: #f9fafb; border-radius: 8px;">
                                        <h3 style="margin: 0 0 16px 0; color: #1a2332; font-size: 16px; font-weight: 600;">✨ Wat kun je met Lexi?</h3>
                                        <ul style="margin: 0; padding-left: 20px; color: #374151; font-size: 14px; line-height: 2;">
                                            <li>Stel CAO-vragen en krijg direct antwoorden</li>
                                            <li>Gebaseerd op 1.000+ officiële documenten</li>
                                            <li>Upload eigen documenten voor analyse</li>
                                            <li>Genereer contracten en brieven</li>
                                        </ul>
                                    </div>
                                    <div style="background-color: #fef2f2; border-left: 4px solid #DC2626; border-radius: 8px; padding: 16px; margin: 24px 0;">
                                        <p style="margin: 0; color: #991b1b; font-size: 13px; line-height: 1.6;">
                                            <strong>⚠️ Veiligheid:</strong> Deze activatie link is 24 uur geldig en kan maar één keer worden gebruikt.
                                        </p>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong> - Jouw Expert CAO Assistent
                                    </p>
                                    <p style="margin: 0; color: #9ca3af; font-size: 12px;">
                                        Vragen? support@lexiai.nl
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        
        return self.send_email(user.email, subject, html_content)
    
    def send_password_reset_link_email(self, user, tenant, reset_url):
        """Send password reset link email (NO password in email - token-based)"""
        subject = "Wachtwoord resetten - Lexi CAO Meester"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <!-- Header -->
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">
                                        LEXI
                                    </h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">
                                        CAO MEESTER
                                    </p>
                                </td>
                            </tr>
                            
                            <!-- Content -->
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">
                                        🔒 Wachtwoord Reset Aangevraagd
                                    </h2>
                                    
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Hoi {user.first_name},
                                    </p>
                                    
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        We hebben een verzoek ontvangen om je wachtwoord te resetten voor je Lexi CAO Meester account bij <strong style="color: #1a2332;">{tenant.company_name}</strong>.
                                    </p>
                                    
                                    <!-- Reset Link Box -->
                                    <div style="background-color: #f9fafb; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <h3 style="margin: 0 0 16px 0; color: #1a2332; font-size: 18px; font-weight: 600;">
                                            🔑 Reset je wachtwoord
                                        </h3>
                                        
                                        <p style="margin: 0 0 16px 0; color: #6b7280; font-size: 14px; line-height: 1.5;">
                                            Klik op de onderstaande knop om een nieuw wachtwoord in te stellen. Deze link is <strong>1 uur geldig</strong> en kan maar <strong>één keer gebruikt</strong> worden.
                                        </p>
                                        
                                        <!-- CTA Button -->
                                        <table width="100%" cellpadding="0" cellspacing="0" style="margin: 16px 0;">
                                            <tr>
                                                <td align="center">
                                                    <a href="{reset_url}" 
                                                       style="display: inline-block; background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); color: #d4af37; text-decoration: none; padding: 16px 48px; border-radius: 8px; font-size: 16px; font-weight: 600; letter-spacing: 0.5px; box-shadow: 0 4px 12px rgba(26, 35, 50, 0.3);">
                                                        Wachtwoord Resetten →
                                                    </a>
                                                </td>
                                            </tr>
                                        </table>
                                        
                                        <p style="margin: 16px 0 0 0; color: #9ca3af; font-size: 12px; line-height: 1.5; word-break: break-all;">
                                            Of kopieer deze link: <br>
                                            <span style="color: #6b7280;">{reset_url}</span>
                                        </p>
                                    </div>
                                    
                                    <!-- Security Notice -->
                                    <div style="margin: 32px 0; padding: 20px; background-color: #fef3c7; border-left: 4px solid #f59e0b; border-radius: 8px;">
                                        <p style="margin: 0; color: #92400e; font-size: 14px; line-height: 1.6;">
                                            <strong>⚡ Veiligheidswaarschuwing:</strong> Heb je deze wachtwoordreset NIET aangevraagd? Negeer deze email en je account blijft veilig. Neem contact op met je administrator als je dit verdacht vindt.
                                        </p>
                                    </div>
                                    
                                    <div style="margin: 24px 0; padding: 16px; background-color: #eff6ff; border-left: 4px solid #3b82f6; border-radius: 8px;">
                                        <p style="margin: 0; color: #1e40af; font-size: 13px; line-height: 1.6;">
                                            💡 <strong>Tip:</strong> Deze link werkt maar 1 keer en verloopt over 1 uur. Als de link niet meer werkt, kun je een nieuwe aanvragen.
                                        </p>
                                    </div>
                                </td>
                            </tr>
                            
                            <!-- Footer -->
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        Veilig wachtwoord resetten! 🔐
                                    </p>
                                    <p style="margin: 0; color: #9ca3af; font-size: 13px;">
                                        Het <strong style="color: #d4af37;">Lexi AI</strong> Team
                                    </p>
                                    
                                    <p style="margin: 24px 0 0 0; color: #9ca3af; font-size: 12px; line-height: 1.6;">
                                        Deze email is verstuurd naar {user.email}<br>
                                        omdat er een wachtwoordreset is aangevraagd.
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        
        return self.send_email(user.email, subject, html_content)
    
    def send_password_reset_email(self, user, tenant, new_password, login_url):
        """Send password reset email with new credentials"""
        subject = "Je wachtwoord is gereset - Lexi CAO Meester"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <!-- Header -->
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">
                                        LEXI
                                    </h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">
                                        CAO MEESTER
                                    </p>
                                </td>
                            </tr>
                            
                            <!-- Content -->
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">
                                        🔒 Je wachtwoord is gereset
                                    </h2>
                                    
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Hoi {user.first_name},
                                    </p>
                                    
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Je wachtwoord voor je Lexi CAO Meester account bij <strong style="color: #1a2332;">{tenant.company_name}</strong> is gereset.
                                    </p>
                                    
                                    <!-- New Password Box -->
                                    <div style="background-color: #f9fafb; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <h3 style="margin: 0 0 16px 0; color: #1a2332; font-size: 18px; font-weight: 600;">
                                            🔑 Je nieuwe wachtwoord
                                        </h3>
                                        
                                        <div style="background-color: #ffffff; padding: 16px; border-radius: 4px; border: 2px solid #d4af37; margin: 16px 0;">
                                            <p style="margin: 0; color: #6b7280; font-size: 13px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">
                                                Nieuw Wachtwoord:
                                            </p>
                                            <p style="margin: 8px 0 0 0; color: #1a2332; font-size: 20px; font-family: 'Courier New', monospace; font-weight: 700; letter-spacing: 1px;">
                                                {new_password}
                                            </p>
                                        </div>
                                        
                                        <p style="margin: 16px 0 0 0; color: #6b7280; font-size: 13px; line-height: 1.5;">
                                            ⚠️ <strong>Belangrijk:</strong> Wijzig dit wachtwoord direct na inloggen via je profielinstellingen voor extra veiligheid.
                                        </p>
                                    </div>
                                    
                                    <!-- CTA Button -->
                                    <table width="100%" cellpadding="0" cellspacing="0" style="margin: 32px 0;">
                                        <tr>
                                            <td align="center">
                                                <a href="{login_url}" 
                                                   style="display: inline-block; background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); color: #d4af37; text-decoration: none; padding: 16px 48px; border-radius: 8px; font-size: 16px; font-weight: 600; letter-spacing: 0.5px; box-shadow: 0 4px 12px rgba(26, 35, 50, 0.3);">
                                                    Inloggen →
                                                </a>
                                            </td>
                                        </tr>
                                    </table>
                                    
                                    <!-- Security Notice -->
                                    <div style="margin: 32px 0; padding: 20px; background-color: #fef3c7; border-left: 4px solid #f59e0b; border-radius: 8px;">
                                        <p style="margin: 0; color: #92400e; font-size: 14px; line-height: 1.6;">
                                            <strong>⚡ Veiligheids tip:</strong> Heb je deze wachtwoordreset niet aangevraagd? Neem dan direct contact op met je administrator.
                                        </p>
                                    </div>
                                </td>
                            </tr>
                            
                            <!-- Footer -->
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        Veilig inloggen! 🔐
                                    </p>
                                    <p style="margin: 0; color: #9ca3af; font-size: 13px;">
                                        Het <strong style="color: #d4af37;">Lexi AI</strong> Team
                                    </p>
                                    
                                    <p style="margin: 24px 0 0 0; color: #9ca3af; font-size: 12px; line-height: 1.6;">
                                        Deze email is verstuurd naar {user.email}<br>
                                        omdat je wachtwoord is gereset.
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        
        return self.send_email(user.email, subject, html_content)
    
    def send_payment_failed_email(self, tenant):
        subject = "⚠️ Betaling mislukt"
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <h2>⚠️ Betaling mislukt</h2>
            <p>Hoi {tenant.contact_name},</p>
            <p>We konden je laatste betaling voor Lexi CAO Meester niet verwerken.</p>
            <p>Update je betaalmethode om actief te blijven en toegang te behouden tot Lexi.</p>
            <p><a href="https://lexiai.nl/admin/billing" 
               style="background: #DC2626; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
               Betaalmethode Updaten →
            </a></p>
            <br>
            <p>Groeten,<br>Het Lexi team</p>
        </body>
        </html>
        """
        return self.send_email(tenant.contact_email, subject, html_content)
    
    def send_trial_expiring_email(self, tenant, days_left):
        subject = f"⏰ Je trial verloopt over {days_left} dagen"
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <h2>⏰ Je trial verloopt binnenkort</h2>
            <p>Hoi {tenant.contact_name},</p>
            <p>Je 14-daagse trial van Lexi CAO Meester verloopt over {days_left} dagen.</p>
            <p>Upgrade nu naar een betaald plan om toegang te behouden tot Lexi en al je chat geschiedenis.</p>
            <p><strong>Beschikbare plannen:</strong></p>
            <ul>
                <li>Professional: €499/maand (5 users, unlimited questions)</li>
                <li>Enterprise: €1.199/maand (unlimited users, unlimited questions)</li>
            </ul>
            <p><a href="https://lexiai.nl/admin/billing" 
               style="background: #4F46E5; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block;">
               Upgrade Nu →
            </a></p>
            <br>
            <p>Groeten,<br>Het Lexi team</p>
        </body>
        </html>
        """
        return self.send_email(tenant.contact_email, subject, html_content)
    
    def send_payment_success_email(self, tenant, plan, amount):
        """Send email after successful payment/subscription activation"""
        subject = f"✅ Welkom bij Lexi CAO Meester - {plan.title()} Plan Actief!"
        
        plan_details = {
            'starter': ('Starter', '€499', '3 users'),
            'professional': ('Professional', '€599', '5 users'),
            'enterprise': ('Enterprise', '€1.199', 'Unlimited users')
        }
        
        plan_name, plan_price, plan_users = plan_details.get(plan, ('Professional', '€599', '5 users'))
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Betaling Succesvol! 🎉</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {tenant.contact_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Bedankt voor je betaling! Je <strong>{plan_name}</strong> abonnement is nu actief voor <strong>{tenant.company_name}</strong>.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <h3 style="margin: 0 0 16px 0; color: #1a2332; font-size: 18px; font-weight: 600;">📋 Abonnement Details</h3>
                                        <ul style="margin: 0; padding-left: 20px; color: #374151; line-height: 1.8;">
                                            <li><strong>Plan:</strong> {plan_name}</li>
                                            <li><strong>Prijs:</strong> {plan_price}/maand</li>
                                            <li><strong>Gebruikers:</strong> {plan_users}</li>
                                            <li><strong>Bedrijf:</strong> {tenant.company_name}</li>
                                        </ul>
                                    </div>
                                    <div style="background-color: #d4af37; border-radius: 8px; padding: 20px; margin: 24px 0; text-align: center;">
                                        <p style="margin: 0 0 12px 0; color: #1a2332; font-size: 18px; font-weight: 600;">Start nu met Lexi!</p>
                                        <a href="https://lexiai.nl/chat" style="background: #1a2332; color: #d4af37; padding: 12px 32px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: 600;">
                                            Naar Chat →
                                        </a>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong> - Jouw Expert CAO Assistent
                                    </p>
                                    <p style="margin: 0; color: #9ca3af; font-size: 12px;">
                                        Vragen? Neem contact op via support@lexiai.nl
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(tenant.contact_email, subject, html_content)
    
    def send_subscription_updated_email(self, tenant, old_plan, new_plan):
        """Send email when subscription plan changes"""
        subject = f"✅ Je abonnement is gewijzigd naar {new_plan.title()}"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Abonnement Gewijzigd</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {tenant.contact_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Je abonnement voor <strong>{tenant.company_name}</strong> is gewijzigd.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">Oud plan:</p>
                                        <p style="margin: 0 0 16px 0; color: #1a2332; font-size: 18px; font-weight: 600;">{old_plan.title()}</p>
                                        <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">Nieuw plan:</p>
                                        <p style="margin: 0; color: #d4af37; font-size: 18px; font-weight: 600;">{new_plan.title()}</p>
                                    </div>
                                    <p style="margin: 0; color: #6b7280; font-size: 14px; text-align: center;">
                                        De wijziging is direct actief.
                                    </p>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong>
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(tenant.contact_email, subject, html_content)
    
    def send_subscription_cancelled_email(self, tenant):
        """Send email when subscription is cancelled"""
        subject = "Je abonnement is geannuleerd"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Abonnement Geannuleerd</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {tenant.contact_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Je abonnement voor <strong>{tenant.company_name}</strong> is geannuleerd.
                                    </p>
                                    <div style="background-color: #fef2f2; border-left: 4px solid #DC2626; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <p style="margin: 0 0 12px 0; color: #1a2332; font-size: 16px; font-weight: 600;">Wat betekent dit?</p>
                                        <ul style="margin: 0; padding-left: 20px; color: #374151; line-height: 1.8;">
                                            <li>Je toegang blijft actief tot het einde van je huidige factuurperiode</li>
                                            <li>Daarna wordt je account gedeactiveerd</li>
                                            <li>Al je chat geschiedenis blijft bewaard</li>
                                        </ul>
                                    </div>
                                    <p style="margin: 24px 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6; text-align: center;">
                                        Mocht je van gedachten veranderen, je bent altijd welkom terug!
                                    </p>
                                    <div style="text-align: center;">
                                        <a href="https://lexiai.nl/admin/billing" style="background: #d4af37; color: #1a2332; padding: 12px 32px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: 600;">
                                            Heractiveer Abonnement
                                        </a>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0; color: #6b7280; font-size: 14px;">
                                        Vragen? support@lexiai.nl
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(tenant.contact_email, subject, html_content)
    
    def send_ideal_payment_link_email(self, user, tenant, invoice_url, amount, due_date):
        """Send monthly iDEAL payment link email"""
        subject = f"💳 Maandelijkse betaling - {amount} via iDEAL"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">💳 Maandelijkse betaling via iDEAL</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {user.first_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Het is tijd voor je maandelijkse betaling van <strong>{amount}</strong> voor {tenant.company_name}.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <p style="margin: 0 0 12px 0; color: #1a2332; font-size: 16px; font-weight: 600;">Factuurdetails:</p>
                                        <div style="color: #374151; line-height: 1.8;">
                                            <p style="margin: 0 0 8px 0;"><strong>Bedrag:</strong> {amount}</p>
                                            <p style="margin: 0 0 8px 0;"><strong>Vervaldatum:</strong> {due_date}</p>
                                            <p style="margin: 0;"><strong>Betaalmethode:</strong> iDEAL</p>
                                        </div>
                                    </div>
                                    <div style="background-color: #fef9f3; border-radius: 8px; padding: 20px; margin: 24px 0;">
                                        <p style="margin: 0 0 12px 0; color: #92400e; font-size: 14px; line-height: 1.6;">
                                            ℹ️ <strong>Let op:</strong> Je hebt gekozen voor handmatige betaling via iDEAL. Klik op de knop hieronder om je betaling te voltooien via je bank.
                                        </p>
                                        <p style="margin: 0; color: #92400e; font-size: 14px; line-height: 1.6;">
                                            <strong>Wil je automatische incasso?</strong> Stuur een email naar <a href="mailto:support@lexiai.nl" style="color: #d4af37;">support@lexiai.nl</a> voor SEPA automatische incasso (over 30 dagen beschikbaar).
                                        </p>
                                    </div>
                                    <div style="text-align: center; margin: 32px 0;">
                                        <a href="{invoice_url}" style="background: #d4af37; color: #1a2332; padding: 16px 48px; text-decoration: none; border-radius: 8px; display: inline-block; font-weight: 600; font-size: 18px; box-shadow: 0 4px 6px rgba(212, 175, 55, 0.3);">
                                            Betaal nu via iDEAL
                                        </a>
                                    </div>
                                    <p style="margin: 24px 0 0 0; color: #6b7280; font-size: 14px; line-height: 1.6; text-align: center;">
                                        Deze betaallink is 30 dagen geldig. Betaal op tijd om toegang te behouden.
                                    </p>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">
                                        Vragen over je factuur?
                                    </p>
                                    <p style="margin: 0; color: #6b7280; font-size: 14px;">
                                        <a href="mailto:support@lexiai.nl" style="color: #d4af37; text-decoration: none;">support@lexiai.nl</a>
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(user.email, subject, html_content)
    
    def send_role_changed_email(self, user, tenant, new_role, changed_by):
        """Send email when user role is changed"""
        subject = f"Je rol is gewijzigd in Lexi CAO Meester"
        
        role_names = {
            'USER': 'Gebruiker',
            'TENANT_ADMIN': 'Administrator'
        }
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Je Rol is Gewijzigd</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {user.first_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        {changed_by} heeft je rol gewijzigd in <strong>{tenant.company_name}</strong>.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0; text-align: center;">
                                        <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">Je nieuwe rol:</p>
                                        <p style="margin: 0; color: #1a2332; font-size: 24px; font-weight: 600;">{role_names.get(new_role, new_role)}</p>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong>
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(user.email, subject, html_content)
    
    def send_account_deactivated_email(self, user, tenant, deactivated_by):
        """Send email when user account is deactivated"""
        subject = "Je account is gedeactiveerd"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">Account Gedeactiveerd</h2>
                                    <p style="margin: 0 0 16px 0; color: #374151; font-size: 16px; line-height: 1.6;">Hoi {user.first_name},</p>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        {deactivated_by} heeft je account gedeactiveerd bij <strong>{tenant.company_name}</strong>.
                                    </p>
                                    <div style="background-color: #fef2f2; border-left: 4px solid #DC2626; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <p style="margin: 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                            Je hebt geen toegang meer tot Lexi CAO Meester. Neem contact op met je administrator voor meer informatie.
                                        </p>
                                    </div>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI</strong>
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(user.email, subject, html_content)
    
    def send_ticket_resolved_email(self, ticket, tenant):
        """Send email when support ticket is resolved"""
        subject = f"Support ticket #{ticket.id} opgelost"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; background-color: #f3f4f6;">
            <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #f3f4f6; padding: 40px 20px;">
                <tr>
                    <td align="center">
                        <table width="600" cellpadding="0" cellspacing="0" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
                            <tr>
                                <td style="background: linear-gradient(135deg, #1a2332 0%, #2a3f5f 100%); padding: 40px 30px; text-align: center;">
                                    <h1 style="margin: 0; color: #d4af37; font-size: 32px; font-weight: 700; letter-spacing: 2px;">LEXI</h1>
                                    <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 14px; letter-spacing: 1px;">CAO MEESTER</p>
                                </td>
                            </tr>
                            <tr>
                                <td style="padding: 40px 30px;">
                                    <h2 style="margin: 0 0 20px 0; color: #1a2332; font-size: 24px; font-weight: 600;">✅ Ticket Opgelost</h2>
                                    <p style="margin: 0 0 24px 0; color: #374151; font-size: 16px; line-height: 1.6;">
                                        Je support ticket is opgelost.
                                    </p>
                                    <div style="background-color: #f0f9ff; border-left: 4px solid #d4af37; border-radius: 8px; padding: 24px; margin: 24px 0;">
                                        <p style="margin: 0 0 8px 0; color: #6b7280; font-size: 14px;">Ticket:</p>
                                        <p style="margin: 0 0 16px 0; color: #1a2332; font-size: 18px; font-weight: 600;">#{ticket.id} - {ticket.subject}</p>
                                    </div>
                                    <p style="margin: 0; color: #374151; font-size: 14px; text-align: center;">
                                        Heb je nog vragen? Open een nieuw ticket via het support menu.
                                    </p>
                                </td>
                            </tr>
                            <tr>
                                <td style="background-color: #f9fafb; padding: 30px; text-align: center; border-top: 1px solid #e5e7eb;">
                                    <p style="margin: 0; color: #6b7280; font-size: 14px;">
                                        <strong style="color: #1a2332;">Lexi AI Support</strong>
                                    </p>
                                </td>
                            </tr>
                        </table>
                    </td>
                </tr>
            </table>
        </body>
        </html>
        """
        return self.send_email(ticket.email, subject, html_content)

# Initialize RAG service (Memgraph + DeepSeek V3)
rag_service = MemgraphDeepSeekService()
s3_service = S3Service()
email_service = EmailService()
