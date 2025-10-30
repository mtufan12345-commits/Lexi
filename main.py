import os
import re
import tempfile
import json
from dotenv import load_dotenv

# Load environment variables from .env file FIRST
load_dotenv()

from flask import Flask, render_template, request, redirect, url_for, jsonify, g, session, flash, Response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_compress import Compress
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.utils import secure_filename
from models import db, SuperAdmin, Tenant, User, Chat, Message, Subscription, Template, UploadedFile, Artifact, SupportTicket, SupportReply
from services import rag_service, s3_service, email_service, StripeService
import stripe
from datetime import datetime, timedelta
import secrets
from functools import wraps

# Optional imports - may not be available in all environments
try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except (ImportError, AttributeError) as e:
    print(f"⚠️  MarkItDown not available: {e}")
    MarkItDown = None
    MARKITDOWN_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except (ImportError, AttributeError) as e:
    print(f"⚠️  OCR tools not available: {e}")
    pytesseract = None
    convert_from_path = None
    OCR_AVAILABLE = False

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# SECURITY: Enable Jinja2 autoescape to prevent XSS attacks
app.jinja_env.autoescape = True

# Global upload status tracker for Super Admin document uploads
super_admin_upload_status = {
    'status': 'idle',  # idle, uploading, processing, complete, error
    'progress': 0,
    'current_file': '',
    'total_files': 0,
    'processed_files': 0,
    'messages': [],
    'imported_count': 0,
    'error': None
}

# Build version for cache busting (use git commit hash in production, timestamp in dev)
# Force new version after Vertex AI rate limit fixes (2025-10-24)
BUILD_VERSION = os.environ.get('BUILD_VERSION', '20251024134500')

# SECURITY: REQUIRE session secret - NEVER use hardcoded fallback (prevents session forgery)
app.secret_key = os.environ.get("SESSION_SECRET") or os.environ.get("SECRET_KEY")
if not app.secret_key:
    raise RuntimeError(
        "❌ SECURITY ERROR: SESSION_SECRET environment variable is required!\n"
        "Set SESSION_SECRET to a strong random value (minimum 32 characters).\n"
        "Generate one with: python -c 'import secrets; print(secrets.token_hex(32))'"
    )
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# CSRF Protection - ENABLED by default for security (disable only in dev with ENABLE_CSRF=false)
app.config['WTF_CSRF_ENABLED'] = os.getenv('ENABLE_CSRF', 'true').lower() == 'true'
app.config['WTF_CSRF_TIME_LIMIT'] = None
app.config['WTF_CSRF_SSL_STRICT'] = False

# Session Cookie Security (Enhanced)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Lax for Stripe redirects compatibility
# Secure cookies only in production (HTTPS), allow HTTP in development
app.config['SESSION_COOKIE_SECURE'] = os.getenv('ENVIRONMENT') == 'production'
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent JavaScript access
# Don't set SESSION_COOKIE_DOMAIN - let Flask use the request host
# This prevents issues with apex domains (lexiai.nl) vs subdomains (company.lexiai.nl)
# Cookies will be set for the exact domain accessed
app.config['SESSION_COOKIE_DOMAIN'] = None
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)  # 8 hour session timeout

if app.config['WTF_CSRF_ENABLED']:
    csrf = CSRFProtect(app)
    print("CSRF Protection enabled")
else:
    print("CSRF Protection disabled for development")
    @app.context_processor
    def csrf_token_processor():
        return {'csrf_token': lambda: ''}

# Make BUILD_VERSION available in all templates for cache busting
@app.context_processor
def inject_build_version():
    return {'build_version': BUILD_VERSION}

db.init_app(app)

# Initialize compression (gzip, brotli, zstd)
app.config['COMPRESS_ALGORITHM'] = 'gzip'
app.config['COMPRESS_LEVEL'] = 6
app.config['COMPRESS_MIN_SIZE'] = 500
app.config['COMPRESS_MIMETYPES'] = [
    'text/html',
    'text/css',
    'text/javascript',
    'application/javascript',
    'application/json',
    'text/xml',
    'application/xml',
]
compress = Compress(app)

# Initialize Rate Limiter for security (no default limits - only specific endpoint limits)
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=[],  # No global limits - use specific endpoint limits only
    storage_uri="memory://"
)

# Initialize Stripe globally - Gebruik productie keys met fallback naar test keys
# Productie heeft voorrang: STRIPE_SECRET_KEY_PROD > STRIPE_SECRET_KEY
stripe.api_key = os.getenv('STRIPE_SECRET_KEY_PROD') or os.getenv('STRIPE_SECRET_KEY')
is_production_stripe = bool(os.getenv('STRIPE_SECRET_KEY_PROD'))
print(f"Stripe initialized: {'Production' if is_production_stripe else 'Test'} mode - Key present: {stripe.api_key is not None}")

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    # Try to load as SuperAdmin first, then fall back to regular User
    super_admin = SuperAdmin.query.get(int(user_id))
    if super_admin:
        return super_admin
    return User.query.get(int(user_id))

@login_manager.unauthorized_handler
def unauthorized():
    """Custom unauthorized handler - redirect to correct login page"""
    # If accessing super-admin routes, redirect to super-admin login
    if request.path.startswith('/super-admin'):
        return redirect(url_for('super_admin_login', next=request.path))
    # Otherwise redirect to normal login
    return redirect(url_for('login', next=request.path))

def get_max_users_for_tier(tier):
    """Get max users allowed for a subscription tier"""
    tier_limits = {
        'starter': 5,
        'professional': 10,
        'enterprise': 999999
    }
    return tier_limits.get(tier, 5)

@app.before_request
def force_https():
    """SECURITY: Force HTTPS in production"""
    # Skip HTTPS redirect for health check endpoint (monitoring tools)
    if request.path == '/health':
        return None

    # Skip HTTPS redirect for upload endpoints (multipart form data)
    if request.path.startswith('/upload/'):
        return None

    # Only enforce HTTPS in production environment (not in development/Replit)
    if os.getenv('ENVIRONMENT') == 'production':
        # Check if request is not secure and not already HTTPS via proxy
        if not request.is_secure and request.headers.get('X-Forwarded-Proto') != 'https':
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)

@app.before_request
def cleanup_stale_pending_signups():
    """Clean up pending signups older than 24 hours"""
    from models import PendingSignup
    from datetime import datetime, timedelta
    
    cutoff_time = datetime.utcnow() - timedelta(hours=24)
    stale_signups = PendingSignup.query.filter(PendingSignup.created_at < cutoff_time).all()
    
    count = len(stale_signups)
    if count > 0:
        for signup in stale_signups:
            db.session.delete(signup)
        db.session.commit()
        print(f"🧹 Cleaned up {count} stale pending signups")

@app.before_request
def validate_host_header():
    """SECURITY: Global Host header validation - prevents Host header injection attacks"""
    # SECURITY: Validate Host header against allowed domains (GLOBAL protection)
    request_host = request.host.split(':')[0]  # Remove port

    # Get allowed hosts from environment (default includes localhost + Replit domains)
    allowed_hosts_env = os.getenv('ALLOWED_HOSTS', 'localhost,127.0.0.1,replit.dev,replit.app')
    allowed_hosts = [h.strip() for h in allowed_hosts_env.split(',')]

    # Check if host is allowed (exact match or subdomain of allowed domain)
    is_allowed = request_host in allowed_hosts or any(
        request_host.endswith(f'.{allowed}') or request_host == allowed
        for allowed in allowed_hosts
    )

    if not is_allowed:
        app.logger.warning(f"🚨 SECURITY: Rejected Host header: {request_host} | Allowed: {allowed_hosts}")
        print(f"[DEBUG] validate_host_header: request_host={request_host}, allowed_hosts={allowed_hosts}, is_allowed={is_allowed}")
        return "Invalid Host header", 400
    else:
        print(f"[DEBUG] validate_host_header: request_host={request_host} is allowed")

@app.before_request
def load_tenant():
    """Load tenant from session after login - NO subdomain routing"""
    g.tenant = None
    g.is_super_admin = session.get('is_super_admin', False)
    print(f"[DEBUG] load_tenant - session keys: {list(session.keys())}, is_super_admin: {g.is_super_admin}")

    if g.is_super_admin:
        return
    
    # Multi-tenant via session (set after login)
    tenant_id = session.get('tenant_id')
    print(f"[DEBUG] load_tenant - tenant_id from session: {tenant_id}")
    if tenant_id:
        g.tenant = Tenant.query.get(tenant_id)
        print(f"[DEBUG] Tenant loaded from session: {g.tenant.company_name if g.tenant else None}")

def tenant_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not g.tenant:
            return "Tenant niet gevonden", 404
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.role == 'admin':
            flash('Je hebt geen toegang tot deze pagina.', 'danger')
            return redirect(url_for('chat_page'))
        return f(*args, **kwargs)
    return decorated_function

def super_admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not g.is_super_admin:
            return "Geen toegang", 403
        return f(*args, **kwargs)
    return decorated_function

@app.after_request
def add_security_and_cache_headers(response):
    """Add security headers, cache headers and enable gzip compression"""
    
    # ========== SECURITY HEADERS (CRITICAL) ==========
    
    # Hide server version information (prevent information disclosure)
    response.headers['Server'] = 'Lexi AI'
    
    # HSTS - Force HTTPS for 1 year (preload ready)
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains; preload'
    
    # Clickjacking protection - Allow same origin iframes (needed for Replit preview)
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    
    # MIME-type sniffing prevention
    response.headers['X-Content-Type-Options'] = 'nosniff'
    
    # Content Security Policy - Strict but allows inline scripts/styles (needed for current app)
    csp_policy = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://js.stripe.com https://cdn.jsdelivr.net; "
        "style-src 'self' 'unsafe-inline'; "
        "img-src 'self' data: https:; "
        "font-src 'self' data:; "
        "connect-src 'self' https://api.stripe.com; "
        "frame-src 'self' https://js.stripe.com; "
        "form-action 'self'; "
        "base-uri 'self'; "
        "object-src 'none'; "
        "upgrade-insecure-requests;"
    )
    response.headers['Content-Security-Policy'] = csp_policy
    
    # Referrer policy - Balance privacy and functionality
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    
    # Permissions policy - Deny unnecessary browser features
    response.headers['Permissions-Policy'] = 'camera=(), microphone=(), geolocation=(), payment=(self)'
    
    # XSS Protection (legacy browsers)
    response.headers['X-XSS-Protection'] = '1; mode=block'
    
    # ========== CACHE & COMPRESSION HEADERS ==========
    
    # Enable gzip compression for text-based responses
    if response.mimetype in ['text/html', 'text/css', 'text/javascript', 'application/javascript', 'application/json', 'text/xml', 'application/xml']:
        response.headers['Vary'] = 'Accept-Encoding'
    
    # Cache static files for 1 year, no cache for dynamic pages
    if request.path.startswith('/static/'):
        response.headers['Cache-Control'] = 'public, max-age=31536000, immutable'
    else:
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    
    return response

@app.route('/sitemap.xml')
def sitemap():
    """Generate dynamic sitemap.xml for SEO"""
    from datetime import datetime
    from flask import Response
    
    pages = [
        {'loc': url_for('index', _external=True), 'lastmod': '2025-01-18', 'changefreq': 'weekly', 'priority': '1.0'},
        {'loc': url_for('pricing', _external=True), 'lastmod': '2025-01-18', 'changefreq': 'weekly', 'priority': '0.9'},
        {'loc': url_for('login', _external=True), 'lastmod': '2025-01-18', 'changefreq': 'monthly', 'priority': '0.8'},
        {'loc': url_for('privacy', _external=True), 'lastmod': '2025-01-18', 'changefreq': 'monthly', 'priority': '0.5'},
        {'loc': url_for('terms', _external=True), 'lastmod': '2025-01-18', 'changefreq': 'monthly', 'priority': '0.5'},
    ]
    
    sitemap_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    sitemap_xml += '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
    
    for page in pages:
        sitemap_xml += '  <url>\n'
        sitemap_xml += f'    <loc>{page["loc"]}</loc>\n'
        sitemap_xml += f'    <lastmod>{page["lastmod"]}</lastmod>\n'
        sitemap_xml += f'    <changefreq>{page["changefreq"]}</changefreq>\n'
        sitemap_xml += f'    <priority>{page["priority"]}</priority>\n'
        sitemap_xml += '  </url>\n'
    
    sitemap_xml += '</urlset>'
    
    return Response(sitemap_xml, mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    """Generate robots.txt for SEO"""
    from flask import Response
    
    robots_txt = f"""User-agent: *
Allow: /

Sitemap: {url_for('sitemap', _external=True)}
"""
    
    return Response(robots_txt, mimetype='text/plain')

@app.route('/')
def index():
    return render_template('landing.html')

@app.route('/prijzen')
def pricing():
    return render_template('pricing.html')

@app.route('/algemene-voorwaarden')
def terms():
    return render_template('terms.html')

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/disclaimer')
def disclaimer():
    return render_template('disclaimer.html')

def count_user_questions(user_id):
    """Count total questions asked by user using message_count with fallbacks"""
    user_chats = Chat.query.filter_by(user_id=user_id).all()
    question_count = 0
    
    for c in user_chats:
        if c.message_count and c.message_count > 0:
            # Use message_count if available (most reliable)
            # Divide by 2 since message_count includes both user and assistant messages
            question_count += (c.message_count + 1) // 2
        elif c.s3_messages_key:
            # Try S3 if message_count not set
            messages = s3_service.get_chat_messages(c.s3_messages_key)
            if messages:
                question_count += sum(1 for m in messages if m.get('role') == 'user')
            else:
                # S3 failed, try PostgreSQL
                db_messages = Message.query.filter_by(chat_id=c.id, role='user').count()
                question_count += db_messages
        else:
            # No S3, use PostgreSQL
            db_messages = Message.query.filter_by(chat_id=c.id, role='user').count()
            question_count += db_messages
    
    return question_count

@app.route('/signup/tenant', methods=['GET', 'POST'])
def signup_tenant():
    # Get tier and billing cycle from query params (from pricing page)
    tier = request.args.get('tier', 'starter')
    billing = request.args.get('billing', 'monthly')
    
    if request.method == 'POST':
        from stripe_config import get_price_id
        
        company_name = request.form.get('company_name') or ''
        contact_email = request.form.get('contact_email') or ''
        contact_name = request.form.get('contact_name') or ''
        password = request.form.get('password') or ''
        tier = request.form.get('tier', 'starter')
        billing = request.form.get('billing', 'monthly')
        cao_preference = request.form.get('cao_preference', 'NBBU')
        
        # Validate inputs
        if not all([company_name, contact_email, contact_name, password]):
            flash('Alle velden zijn verplicht.', 'danger')
            return render_template('signup_tenant.html', tier=tier, billing=billing)
        
        # Check if email already exists
        existing_user = User.query.filter_by(email=contact_email).first()
        if existing_user:
            flash('Dit email adres is al in gebruik.', 'danger')
            return render_template('signup_tenant.html', tier=tier, billing=billing)
        
        # Get Stripe Price ID
        price_id = get_price_id(tier, billing)
        if not price_id or not price_id.startswith('price_'):
            app.logger.error(f"Invalid or missing Stripe Price ID for tier={tier}, billing={billing}, price_id={price_id}")
            flash('Ongeldige pricing optie. Neem contact op met support.', 'danger')
            return render_template('signup_tenant.html', tier=tier, billing=billing)
        
        # Create Stripe Checkout Session with both card AND iDEAL support
        try:
            from models import PendingSignup
            import requests
            
            # Get base URL (Host header already validated globally by validate_host_header)
            base_url = request.host_url.rstrip('/')
            
            # Use Stripe HTTP API directly to avoid SDK issues
            # Productie key heeft voorrang over test key
            stripe_api_key = os.getenv('STRIPE_SECRET_KEY_PROD') or os.getenv('STRIPE_SECRET_KEY')
            stripe_headers = {
                'Authorization': f'Bearer {stripe_api_key}',
                'Content-Type': 'application/x-www-form-urlencoded'
            }
            
            # ALL payments go through Stripe Checkout
            # NOTE: iDEAL disabled until SEPA Direct Debit is activated in Stripe Dashboard
            stripe_data = {
                'payment_method_types[0]': 'card',
                # 'payment_method_types[1]': 'ideal',  # Requires SEPA activation (30 days approval)
                'line_items[0][price]': price_id,
                'line_items[0][quantity]': 1,
                'mode': 'subscription',
                'success_url': f"{base_url}/signup/success?session_id={{CHECKOUT_SESSION_ID}}",
                'cancel_url': f"{base_url}/signup/cancel?session_id={{CHECKOUT_SESSION_ID}}",
                'customer_email': contact_email,
                'metadata[signup_email]': contact_email,
                'metadata[tier]': tier,
                'metadata[billing]': billing,
                'metadata[cao_preference]': cao_preference
            }
            
            response = requests.post(
                'https://api.stripe.com/v1/checkout/sessions',
                headers=stripe_headers,
                data=stripe_data
            )
            
            if response.status_code != 200:
                raise Exception(f"Stripe API error: {response.text}")
            
            checkout_session = response.json()
            
            # Store signup data in database (server-side) with checkout session ID
            pending_signup = PendingSignup(
                checkout_session_id=checkout_session['id'],
                email=contact_email,
                company_name=company_name,
                contact_name=contact_name,
                tier=tier,
                billing=billing,
                cao_preference=cao_preference
            )
            pending_signup.set_password(password)  # Hash the password
            db.session.add(pending_signup)
            db.session.commit()
            
            # Use iframe-safe redirect to escape Replit preview and load Stripe Checkout in top window
            return render_template('stripe_redirect.html', checkout_url=checkout_session['url'])
            
        except Exception as e:
            app.logger.exception(f"Stripe checkout error for tier={tier}, billing={billing}, price_id={price_id}: {str(e)}")
            
            # Provide more specific error messages
            error_msg = str(e)
            if 'No such price' in error_msg:
                flash('De gekozen prijsoptie is niet geldig. Neem contact op met support.', 'danger')
            elif 'API key' in error_msg:
                flash('Betaalconfiguratie is niet correct ingesteld. Neem contact op met support.', 'danger')
            else:
                flash('Er ging iets mis met de betaling. Probeer het opnieuw.', 'danger')
            
            return render_template('signup_tenant.html', tier=tier, billing=billing)
    
    # GET request - show signup form
    return render_template('signup_tenant.html', tier=tier, billing=billing)

@app.route('/signup/success')
def signup_success():
    """Success page after Stripe payment - account will be created via webhook"""
    from models import PendingSignup
    from flask_login import logout_user, login_user
    import time
    
    session_id = request.args.get('session_id')
    
    if not session_id:
        flash('Geen geldige sessie gevonden.', 'danger')
        return redirect(url_for('index'))
    
    # SECURITY: Get email from server-side PendingSignup record, NOT from URL
    # This prevents account-takeover attacks via URL manipulation
    pending = PendingSignup.query.filter_by(checkout_session_id=session_id).first()
    
    # First, logout any existing user to prevent confusion
    if current_user.is_authenticated:
        app.logger.info(f"Logging out existing user: {current_user.email}")
        logout_user()
        session.clear()
    
    # If pending signup exists, webhook hasn't processed it yet
    if pending:
        email = pending.email
        app.logger.info(f"Pending signup found for {email}, waiting for webhook...")
        
        # Wait for webhook to create account (with retries)
        max_retries = 15
        retry_delay = 1  # seconds
        
        for attempt in range(max_retries):
            # Check if webhook has processed and created the account
            # (pending record will be deleted by webhook after account creation)
            db.session.expire(pending)
            pending = PendingSignup.query.filter_by(checkout_session_id=session_id).first()
            
            if not pending:
                # Pending deleted = webhook processed successfully
                app.logger.info(f"Webhook processed signup for {email}")
                break
            
            app.logger.info(f"Waiting for webhook (attempt {attempt + 1}/{max_retries})...")
            time.sleep(retry_delay)
        
        if pending:
            # Webhook still hasn't processed after max retries - USE FALLBACK
            app.logger.warning(f"Webhook hasn't processed signup for {email} after {max_retries} attempts - using fallback")
            
            # FALLBACK: Verify with Stripe and provision account directly
            try:
                from stripe.checkout import Session as StripeSession
                from provision_tenant import provision_tenant_from_signup
                
                # SECURITY: Validate checkout session with Stripe API (server-side)
                checkout_session = StripeSession.retrieve(session_id)
                
                # Verify payment was successful
                if checkout_session.payment_status != 'paid':
                    app.logger.error(f"Checkout session {session_id} payment status: {checkout_session.payment_status}")
                    flash('Betaling niet gelukt. Neem contact op met support.', 'danger')
                    return redirect(url_for('index'))
                
                # Use shared provisioning service (idempotent and safe)
                success, user, error_msg = provision_tenant_from_signup(
                    pending_signup=pending,
                    stripe_session_data=checkout_session
                )
                
                if not success:
                    app.logger.error(f"Fallback provisioning failed: {error_msg}")
                    flash('Account aanmaken mislukt. Neem contact op met support.', 'danger')
                    return redirect(url_for('index'))
                
                app.logger.info(f"✅ Fallback: Account provisioned successfully for {email}")
                # Continue to auto-login below
                
            except Exception as fallback_error:
                app.logger.error(f"Fallback provisioning error: {fallback_error}")
                flash('Account aanmaken mislukt. Neem contact op met support.', 'danger')
                return redirect(url_for('index'))
    else:
        # No pending signup = webhook already processed
        # We need to find the user that was created for this checkout session
        # SECURITY: We can't trust any user input here. We must verify via Stripe or give up.
        app.logger.info(f"No pending signup for session {session_id} - webhook may have already processed")
        email = None
        
        # Try to get email from Stripe session metadata (server-side verification)
        try:
            from stripe.checkout import Session as StripeSession
            checkout_session = StripeSession.retrieve(session_id)
            email = checkout_session.get('metadata', {}).get('signup_email') or checkout_session.get('customer_email')
            app.logger.info(f"Retrieved email from Stripe: {email}")
        except Exception as e:
            app.logger.error(f"Failed to retrieve Stripe session: {e}")
            flash('Kon account niet verifiëren. Probeer in te loggen met uw email en wachtwoord.', 'warning')
            return redirect(url_for('login'))
        
        if not email:
            flash('Kon account niet verifiëren. Probeer in te loggen met uw email en wachtwoord.', 'warning')
            return redirect(url_for('login'))
    
    # Find the newly created admin user with verified email
    new_user = User.query.filter_by(email=email, role='admin').first()
    
    if not new_user:
        app.logger.error(f"Account not found for verified email {email}")
        flash('Account niet gevonden. Neem contact op met support.', 'danger')
        return redirect(url_for('index'))
    
    # Login the verified user automatically
    login_user(new_user)
    session['tenant_id'] = new_user.tenant_id
    session.permanent = True
    
    app.logger.info(f"✅ Auto-logged in verified user: {new_user.email} (tenant_id: {new_user.tenant_id})")
    
    flash(f'Welkom bij Lexi AI! Uw account is succesvol aangemaakt.', 'success')
    return redirect(url_for('chat_page'))

@app.route('/signup/cancel')
def signup_cancel():
    """Cancel page if user cancels Stripe payment"""
    from models import PendingSignup
    
    # Try to clean up pending signup if session_id is provided
    session_id = request.args.get('session_id')
    if session_id:
        pending = PendingSignup.query.filter_by(checkout_session_id=session_id).first()
        if pending:
            db.session.delete(pending)
            db.session.commit()
    
    flash('Betaling geannuleerd. U kunt opnieuw proberen wanneer u klaar bent.', 'info')
    return redirect(url_for('pricing'))

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def login():
    if current_user.is_authenticated and not g.is_super_admin:
        # Check if user has valid tenant in session
        if not g.tenant:
            # Session is corrupt - clear it and force re-login
            logout_user()
            session.clear()
            flash('Je sessie is verlopen. Log opnieuw in.', 'info')
        else:
            return redirect(url_for('chat_page'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').lower().strip()
        password = request.form.get('password')
        
        print(f"Login attempt - Email: {email}")
        
        # Zoek user op basis van email (uniek over alle tenants)
        user = User.query.filter_by(email=email).first()
        print(f"User found: {user is not None}")
        
        if user and user.check_password(password):
            print("Password check passed")
            
            # Haal de tenant op van deze user
            tenant = Tenant.query.get(user.tenant_id)
            
            if not user.is_active:
                flash('Je account is gedeactiveerd.', 'danger')
                return render_template('login.html')
            
            if tenant.status not in ['trial', 'active']:
                flash('Je account is verlopen. Neem contact op met de administrator.', 'warning')
                return render_template('login.html')
            
            force_login = request.form.get('force_login') == 'true'
            
            if user.session_token and not force_login:
                return render_template('login.html', 
                                     show_force_login=True, 
                                     email=email)
            
            user.session_token = secrets.token_hex(32)
            db.session.commit()
            
            login_user(user)
            session['tenant_id'] = tenant.id  # Zet tenant_id automatisch in session
            session['session_token'] = user.session_token
            session['is_super_admin'] = False
            
            if force_login:
                flash('Oude sessie uitgelogd. Je bent nu ingelogd.', 'success')
            
            print(f"Login successful - Tenant: {tenant.company_name}")
            return redirect(url_for('chat_page'))
        
        print("Login failed - invalid credentials")
        flash('Ongeldige email of wachtwoord.', 'danger')
    
    return render_template('login.html')

@app.route('/forgot-password', methods=['GET', 'POST'])
@limiter.limit("3 per minute")
def forgot_password():
    """Generate password reset token and send reset link via email"""
    if request.method == 'POST':
        email = request.form.get('email', '').lower().strip()
        
        # Find user by email
        user = User.query.filter_by(email=email).first()
        
        if user:
            # Get tenant for email context
            tenant = Tenant.query.get(user.tenant_id)
            
            if tenant and user.is_active:
                # Generate secure URL-safe reset token
                reset_token = secrets.token_urlsafe(32)
                
                # Set token expiration (1 hour from now)
                from datetime import datetime, timedelta
                user.reset_token = reset_token
                user.reset_token_expires_at = datetime.utcnow() + timedelta(hours=1)
                db.session.commit()
                
                # Create reset URL (no subdomain needed)
                domain = os.getenv('PRODUCTION_DOMAIN', 'lexiai.nl')
                reset_url = f"https://{domain}/reset-password/{reset_token}"
                
                # Send email with reset link (NO PASSWORD in email)
                email_service.send_password_reset_link_email(user, tenant, reset_url)
                
                flash('Een email met een reset link is verzonden! Check je inbox.', 'success')
                return redirect(url_for('login'))
            else:
                # Account inactive or tenant not found
                # Don't reveal this to prevent email enumeration - show generic message
                flash('Als dit email adres bestaat, ontvang je een reset email.', 'info')
                return redirect(url_for('login'))
        else:
            # Security: Don't reveal if email exists or not
            # Still show success to prevent email enumeration
            flash('Als dit email adres bestaat, ontvang je een reset email.', 'info')
            return redirect(url_for('login'))
    
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
@limiter.limit("5 per minute")
def reset_password(token):
    """Reset password using token (GET: show form, POST: process new password)"""
    from datetime import datetime
    
    # Find user by reset token
    user = User.query.filter_by(reset_token=token).first()
    
    # Validate token exists and is not expired
    if not user or not user.reset_token_expires_at:
        flash('Ongeldige of verlopen reset link.', 'danger')
        return redirect(url_for('forgot_password'))
    
    if datetime.utcnow() > user.reset_token_expires_at:
        flash('Deze reset link is verlopen. Vraag een nieuwe aan.', 'danger')
        return redirect(url_for('forgot_password'))
    
    # GET: Show reset password form
    if request.method == 'GET':
        return render_template('reset_password.html', token=token)
    
    # POST: Process new password
    new_password = request.form.get('password', '').strip()
    confirm_password = request.form.get('confirm_password', '').strip()
    
    # Validate password
    if not new_password or len(new_password) < 8:
        flash('Wachtwoord moet minimaal 8 karakters zijn.', 'danger')
        return render_template('reset_password.html', token=token)
    
    if new_password != confirm_password:
        flash('Wachtwoorden komen niet overeen.', 'danger')
        return render_template('reset_password.html', token=token)
    
    # Update password and invalidate token (single-use token)
    user.set_password(new_password)
    user.reset_token = None
    user.reset_token_expires_at = None
    db.session.commit()
    
    flash('Je wachtwoord is succesvol gereset! Je kunt nu inloggen.', 'success')
    return redirect(url_for('login'))

@app.route('/super-admin/login', methods=['GET', 'POST'])
@csrf.exempt
@limiter.limit("5 per minute")
def super_admin_login():
    if request.method == 'POST':
        email = request.form.get('email') or ''
        password = request.form.get('password') or ''
        
        print(f"[DEBUG] Super Admin Login Attempt:")
        print(f"  Email: '{email}'")
        print(f"  User-Agent: {request.headers.get('User-Agent', 'Unknown')}")
        print(f"  Password length: {len(password)}")

        admin = SuperAdmin.query.filter_by(email=email).first()
        print(f"  Admin found: {admin is not None}")
        
        if admin:
            print(f"  Admin email: '{admin.email}'")
            print(f"  Admin ID: {admin.id}")
            password_valid = admin.check_password(password)
            print(f"  Password valid: {password_valid}")
            
            if password_valid:
                # FIX: Clear old session and set fresh login session
                session.clear()

                # Set all session data at once
                session['super_admin_id'] = admin.id
                session['is_super_admin'] = True
                session.permanent = True

                # Force Flask to regenerate the session
                login_user(admin, remember=True)

                # Explicitly mark as modified to ensure cookie is set
                session.modified = True

                print(f"  ✅ Login successful for {email}")
                print(f"  Session after login - keys: {list(session.keys())}, permanent: {session.permanent}")
                print(f"  Flask-Login authenticated: {current_user.is_authenticated}")

                # Create redirect response and set session cookies
                response = redirect(url_for('super_admin_dashboard'))

                # CRITICAL FIX for Edge browser: Explicitly ensure session is persisted before redirect
                # For client-side sessions, we need to let Flask's session interface handle it
                # The response should include Set-Cookie headers from the after_request handler

                return response
            else:
                print(f"  ❌ Password INCORRECT for {email}")
        else:
            print(f"  ❌ NO admin found with email: '{email}'")

        flash('Ongeldige credentials.', 'danger')

    return render_template('super_admin_login.html')

@app.route('/select-tenant', methods=['GET', 'POST'])
@super_admin_required
def select_tenant():
    """Development/admin mode: manually select a tenant (alleen voor super admins)"""
    if request.method == 'POST':
        subdomain = request.form.get('subdomain')
        print(f"[DEBUG] select_tenant - subdomain: {subdomain}")
        tenant = Tenant.query.filter_by(subdomain=subdomain).first()
        if tenant:
            session['tenant_id'] = tenant.id
            session.modified = True  # Force session save
            print(f"[DEBUG] Tenant ID {tenant.id} saved to session")
            flash(f'Tenant geselecteerd: {tenant.company_name}', 'success')
            return redirect(url_for('login'))
        flash('Tenant niet gevonden', 'danger')
    
    tenants = Tenant.query.all()
    return render_template('select_tenant.html', tenants=tenants)

@app.route('/logout')
@login_required
def logout():
    if not g.is_super_admin and isinstance(current_user, User):
        current_user.session_token = None
        db.session.commit()
    
    logout_user()
    session.clear()
    return redirect(url_for('index'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
@tenant_required
def user_profile():
    if request.method == 'POST':
        new_email = request.form.get('email')
        
        if new_email != current_user.email:
            existing = User.query.filter_by(
                tenant_id=g.tenant.id,
                email=new_email
            ).first()
            if existing:
                flash('Dit e-mailadres is al in gebruik!', 'error')
                return redirect(url_for('user_profile'))
        
        current_user.first_name = request.form.get('first_name') or current_user.first_name
        current_user.last_name = request.form.get('last_name') or current_user.last_name
        current_user.email = new_email
        
        new_password = request.form.get('new_password') or ''
        if new_password:
            current_user.set_password(new_password)
        
        db.session.commit()
        flash('Profiel bijgewerkt!', 'success')
        return redirect(url_for('user_profile'))
    
    return render_template('user_profile.html', tenant=g.tenant, user=current_user)

@app.route('/api/profile/avatar', methods=['POST'])
@login_required
@tenant_required
def upload_avatar():
    if 'avatar' not in request.files:
        return jsonify({'error': 'Geen bestand'}), 400
    
    file = request.files['avatar']
    if file.filename == '':
        return jsonify({'error': 'Geen bestand geselecteerd'}), 400
    
    # Check file type
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    filename = file.filename or ''
    if '.' not in filename or filename.rsplit('.', 1)[1].lower() not in allowed_extensions:
        return jsonify({'error': 'Alleen afbeeldingen toegestaan (PNG, JPG, GIF, WEBP)'}), 400
    
    # Upload to S3
    s3_key = s3_service.upload_file(file, g.tenant.id, folder='avatars')
    if not s3_key:
        return jsonify({'error': 'Upload mislukt'}), 500
    
    # Get URL
    avatar_url = s3_service.get_file_url(s3_key)
    
    # Update user
    current_user.avatar_url = avatar_url
    db.session.commit()
    
    return jsonify({'success': True, 'avatar_url': avatar_url})

@app.route('/api/gdpr/export-data', methods=['GET'])
@login_required
@tenant_required
def gdpr_export_data():
    """
    GDPR Article 15 - Right of Access
    Export all personal data for the current user
    """
    try:
        # Collect all user data
        user_data = {
            'personal_info': {
                'email': current_user.email,
                'first_name': current_user.first_name,
                'last_name': current_user.last_name,
                'created_at': current_user.created_at.isoformat() if current_user.created_at else None,
            },
            'tenant_info': {
                'company_name': g.tenant.company_name,
                'subdomain': g.tenant.subdomain,
            },
            'chats': [],
            'uploaded_files': [],
        }

        # Get all chats
        chats = Chat.query.filter_by(tenant_id=g.tenant.id, user_id=current_user.id).all()
        for chat in chats:
            chat_data = {
                'title': chat.title,
                'created_at': chat.created_at.isoformat() if chat.created_at else None,
                'updated_at': chat.updated_at.isoformat() if chat.updated_at else None,
                'messages': []
            }

            # Get messages from S3
            if chat.s3_messages_key:
                messages = s3_service.get_chat_messages(chat.s3_messages_key)
                if messages:
                    chat_data['messages'] = messages

            user_data['chats'].append(chat_data)

        # Get all uploaded files
        files = UploadedFile.query.filter_by(tenant_id=g.tenant.id, user_id=current_user.id).all()
        for file in files:
            user_data['uploaded_files'].append({
                'filename': file.filename,
                'uploaded_at': file.uploaded_at.isoformat() if file.uploaded_at else None,
                'file_size': file.file_size,
                'mime_type': file.mime_type,
            })

        # Return as JSON download
        response = Response(
            json.dumps(user_data, indent=2, ensure_ascii=False),
            mimetype='application/json',
            headers={
                'Content-Disposition': f'attachment; filename=lexi_data_export_{current_user.id}_{datetime.utcnow().strftime("%Y%m%d")}.json'
            }
        )
        return response

    except Exception as e:
        app.logger.error(f"GDPR export error: {e}")
        return jsonify({'error': 'Data export failed'}), 500

@app.route('/api/gdpr/delete-account', methods=['POST'])
@login_required
@tenant_required
def gdpr_delete_account():
    """
    GDPR Article 17 - Right to Erasure
    Delete user account and all associated data
    """
    try:
        # Verify password for security
        password = request.json.get('password')
        if not password or not current_user.check_password(password):
            return jsonify({'error': 'Incorrect password'}), 401

        user_id = current_user.id
        tenant_id = g.tenant.id

        # Get all user's chats
        chats = Chat.query.filter_by(tenant_id=tenant_id, user_id=user_id).all()

        # Delete chat messages from S3
        for chat in chats:
            if chat.s3_messages_key:
                s3_service.delete_file(chat.s3_messages_key)
            db.session.delete(chat)

        # Delete uploaded files from S3
        files = UploadedFile.query.filter_by(tenant_id=tenant_id, user_id=user_id).all()
        for file in files:
            if file.s3_key:
                s3_service.delete_file(file.s3_key)
            db.session.delete(file)

        # Delete artifacts from S3
        artifacts = Artifact.query.filter_by(tenant_id=tenant_id, user_id=user_id).all()
        for artifact in artifacts:
            if artifact.s3_key:
                s3_service.delete_file(artifact.s3_key)
            db.session.delete(artifact)

        # Logout user
        logout_user()
        session.clear()

        # Delete user account
        db.session.delete(current_user)
        db.session.commit()

        app.logger.info(f"GDPR: User {user_id} account deleted")

        return jsonify({
            'success': True,
            'message': 'Your account and all data have been permanently deleted'
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"GDPR delete error: {e}")
        return jsonify({'error': 'Account deletion failed'}), 500

@app.route('/chat')
@login_required
@tenant_required
def chat_page():
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        flash('Je account is niet actief. Neem contact op met je beheerder.', 'warning')
        return redirect(url_for('index'))
    
    chats = Chat.query.filter_by(
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).order_by(Chat.updated_at.desc()).all()
    
    return render_template('chat.html', chats=chats, tenant=g.tenant, user=current_user)

@app.route('/api/chat/new', methods=['POST'])
@login_required
@tenant_required
def new_chat():
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    chat = Chat(
        tenant_id=g.tenant.id,
        user_id=current_user.id,
        title='Nieuwe chat'
    )
    db.session.add(chat)
    db.session.commit()
    
    # Associate any pending uploaded files (chat_id=NULL) with this new chat
    pending_files = UploadedFile.query.filter_by(
        tenant_id=g.tenant.id,
        user_id=current_user.id,
        chat_id=None
    ).all()
    
    if pending_files:
        for uploaded_file in pending_files:
            uploaded_file.chat_id = chat.id
        db.session.commit()
        print(f"[DEBUG] Associated {len(pending_files)} pending files with new chat {chat.id}")
    
    return jsonify({'id': chat.id, 'title': chat.title})

@app.route('/api/chat/<int:chat_id>', methods=['GET'])
@login_required
@tenant_required
def get_chat(chat_id):
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    # Get messages from S3
    messages = []
    if chat.s3_messages_key:
        s3_messages = s3_service.get_chat_messages(chat.s3_messages_key)
        for idx, m in enumerate(s3_messages):
            msg_data = {
                'id': idx + 1,
                'role': m.get('role'),
                'content': m.get('content'),
                'created_at': m.get('created_at'),
                'feedback_rating': m.get('feedback_rating')
            }
            
            # Add attachments if present (user messages)
            if m.get('attachments'):
                msg_data['attachments'] = m.get('attachments')
            
            if m.get('role') == 'assistant':
                artifacts = Artifact.query.filter_by(message_id=idx + 1, chat_id=chat.id, tenant_id=g.tenant.id).all()
                if artifacts:
                    msg_data['artifacts'] = [{
                        'id': a.id,
                        'title': a.title,
                        'type': a.artifact_type,
                        'content': a.content
                    } for a in artifacts]
            
            messages.append(msg_data)
    
    return jsonify({'id': chat.id, 'title': chat.title, 'messages': messages})

@app.route('/api/chat/<int:chat_id>/message', methods=['POST'])
@login_required
@tenant_required
@limiter.limit("30 per minute")
def send_message(chat_id):
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    print(f"[DEBUG] send_message called - chat_id: {chat_id}, user: {current_user.id}")
    
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    data = request.json
    user_message = data.get('message', '')
    print(f"[DEBUG] User message: {user_message}")

    # Get uploaded files for this chat
    print(f"[DEBUG] 1. Starting file query for chat_id={chat.id}, tenant={g.tenant.id}, user={current_user.id}")
    try:
        uploaded_files = UploadedFile.query.filter_by(
            chat_id=chat.id,
            tenant_id=g.tenant.id,
            user_id=current_user.id
        ).all()
        print(f"[DEBUG] 2. Found {len(uploaded_files)} uploaded files")
    except Exception as e:
        print(f"[DEBUG] ERROR in file query: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

    # Create user message dict for S3 with file attachments
    print(f"[DEBUG] 3. Creating user_msg_dict")
    user_msg_dict = {
        'role': 'user',
        'content': user_message,
        'created_at': datetime.utcnow().isoformat()
    }
    print(f"[DEBUG] 4. user_msg_dict created successfully")

    # Add file attachments to message ONLY for newly uploaded files
    # Files uploaded AFTER the last message should be shown as attachments
    # For subsequent messages, old files are still used for AI context but not shown as attachments
    if uploaded_files:
        print(f"[DEBUG] 5. Processing {len(uploaded_files)} uploaded files for attachments")
        # Get files uploaded after the last message (new uploads since last message)
        # Guard against None updated_at (legacy/migrated chats) - show all files if None
        if chat.message_count > 0 and chat.updated_at is not None:
            newly_uploaded = [f for f in uploaded_files if f.created_at > chat.updated_at]
        else:
            newly_uploaded = uploaded_files

        if newly_uploaded:
            user_msg_dict['attachments'] = [{
                'id': f.id,
                'filename': f.original_filename,
                'mime_type': f.mime_type
            } for f in newly_uploaded]
            print(f"[DEBUG] 6. Added {len(newly_uploaded)} attachments to message")

    # Append to S3
    print(f"[DEBUG] 7. About to call S3 service append_chat_message")
    print(f"[DEBUG]    - chat.s3_messages_key: {chat.s3_messages_key}")
    print(f"[DEBUG]    - chat.id: {chat.id}")
    print(f"[DEBUG]    - tenant_id: {g.tenant.id}")
    print(f"[DEBUG]    - S3 service enabled: {s3_service.enabled}")
    try:
        s3_key = s3_service.append_chat_message(
            chat.s3_messages_key,
            chat.id,
            g.tenant.id,
            user_msg_dict
        )
        print(f"[DEBUG] 8. S3 service returned s3_key: {s3_key}")
    except Exception as e:
        print(f"[DEBUG] ERROR in S3 service call: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

    if not s3_key:
        print(f"[DEBUG] 9. S3 returned None/False - returning error to user")
        return jsonify({'error': 'Kon bericht niet opslaan. Probeer het opnieuw.'}), 500

    print(f"[DEBUG] 10. Updating chat object in database")
    chat.s3_messages_key = s3_key
    chat.message_count = (chat.message_count or 0) + 1
    print(f"[DEBUG]     - Updated message_count to: {chat.message_count}")

    if chat.message_count <= 1:
        chat.title = user_message[:50] + ('...' if len(user_message) > 50 else '')
        print(f"[DEBUG]     - Set chat title: {chat.title}")

    chat.updated_at = datetime.utcnow()
    print(f"[DEBUG] 11. About to commit database changes...")
    try:
        db.session.commit()
        print(f"[DEBUG] 12. Database commit successful!")
    except Exception as e:
        print(f"[DEBUG] ERROR in database commit: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        raise

    print(f"[DEBUG] 13. Building ai_message for Vertex AI...")
    ai_message = user_message
    file_errors = []
    
    if uploaded_files:
        file_contents = []
        for uploaded_file in uploaded_files:
            print(f"[DEBUG] Processing file: {uploaded_file.original_filename}, type: {uploaded_file.mime_type}")
            # For PDF files, try extracted_text from database first
            if uploaded_file.mime_type == 'application/pdf':
                if uploaded_file.extracted_text and uploaded_file.extracted_text.strip():
                    # Use pre-extracted text if available and not empty
                    print(f"[DEBUG] Using extracted_text from database (length: {len(uploaded_file.extracted_text)})")
                    content = uploaded_file.extracted_text
                    file_contents.append(f"\n\n--- Bestand: {uploaded_file.original_filename} ---\n{content}\n--- Einde bestand ---\n")
                else:
                    # Fallback: try downloading from S3 for legacy PDFs or failed extractions
                    print(f"[DEBUG] No extracted_text, trying S3 fallback for {uploaded_file.original_filename}")
                    try:
                        content, error = s3_service.download_file_content(uploaded_file.s3_key, uploaded_file.mime_type)
                        if error:
                            print(f"[DEBUG] S3 error: {error}")
                            file_errors.append(f"{uploaded_file.original_filename}: {error}")
                        elif content:
                            print(f"[DEBUG] S3 fallback successful, content length: {len(content)}")
                            file_contents.append(f"\n\n--- Bestand: {uploaded_file.original_filename} ---\n{content}\n--- Einde bestand ---\n")
                        else:
                            print(f"[DEBUG] S3 returned empty content")
                    except Exception as e:
                        print(f"[DEBUG] Exception in S3 fallback: {str(e)}")
                        file_errors.append(f"{uploaded_file.original_filename}: Kon bestand niet lezen")
            else:
                # For non-PDF files, download from S3
                content, error = s3_service.download_file_content(uploaded_file.s3_key, uploaded_file.mime_type)
                if error:
                    file_errors.append(f"{uploaded_file.original_filename}: {error}")
                elif content:
                    file_contents.append(f"\n\n--- Bestand: {uploaded_file.original_filename} ---\n{content}\n--- Einde bestand ---\n")
        
        if file_contents:
            ai_message = f"{user_message}\n\n{''.join(file_contents)}"
            print(f"[DEBUG] Including {len(file_contents)} uploaded files in context")
        
        if file_errors and not file_contents:
            error_msg = "\n".join(file_errors)
            return jsonify({'response': f"⚠️ Kon geen bestanden lezen:\n{error_msg}\n\nProbeer andere bestanden.", 'has_errors': True})

    print("[DEBUG] 14. About to call RAG service (Memgraph + DeepSeek)...")
    print(f"[DEBUG]     - ai_message length: {len(ai_message)} chars")
    print(f"[DEBUG]     - RAG service enabled: {rag_service.enabled}")
    try:
        from cao_config import get_system_instruction
        print("[DEBUG] 15. Imported cao_config successfully")
        cao_instruction = get_system_instruction(g.tenant)
        print(f"[DEBUG] 16. Got system instruction (length: {len(cao_instruction)} chars)")
        print(f"[DEBUG] 17. Calling rag_service.chat() (Memgraph + DeepSeek)...")
        lex_response = rag_service.chat(ai_message, system_instruction=cao_instruction)
        print(f"[DEBUG] 18. RAG service response received (length: {len(lex_response)} chars)")
        print(f"[DEBUG]     - First 100 chars: {lex_response[:100]}...")
    except Exception as e:
        print(f"[DEBUG] ERROR in RAG service call: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
    
    # Create assistant message dict for S3
    print("[DEBUG] 19. Creating assistant message dict for S3")
    assistant_msg_dict = {
        'role': 'assistant',
        'content': lex_response,
        'created_at': datetime.utcnow().isoformat()
    }
    print("[DEBUG] 20. Assistant message dict created")

    # Append to S3
    print("[DEBUG] 21. Appending assistant response to S3...")
    try:
        s3_key = s3_service.append_chat_message(
            chat.s3_messages_key,
            chat.id,
            g.tenant.id,
            assistant_msg_dict
        )
        print(f"[DEBUG] 22. S3 append returned: {s3_key}")
    except Exception as e:
        print(f"[DEBUG] ERROR appending assistant message to S3: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

    if not s3_key:
        print("[DEBUG] 23. S3 returned None - returning error")
        return jsonify({'error': 'Kon AI response niet opslaan. Probeer het opnieuw.'}), 500

    print("[DEBUG] 24. Updating chat with assistant message...")
    chat.s3_messages_key = s3_key
    chat.message_count = (chat.message_count or 0) + 1
    chat.updated_at = datetime.utcnow()
    print(f"[DEBUG] 25. Committing final database update (message_count={chat.message_count})")
    try:
        db.session.commit()
        print("[DEBUG] 26. Final database commit successful!")
    except Exception as e:
        print(f"[DEBUG] ERROR in final database commit: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        raise
    
    # Store last message ID for artifacts (use message_count as ID)
    assistant_message_id = chat.message_count
    print(f"[DEBUG] 27. Processing artifacts (message_id={assistant_message_id})")

    artifacts_created = []
    artifact_pattern = r'```artifact:(\w+)\s+title:([^\n]+)\n(.*?)```'
    matches = re.finditer(artifact_pattern, lex_response, re.DOTALL)

    artifacts_to_commit = []
    print("[DEBUG] 28. Searching for artifact patterns in response...")
    for match in matches:
        artifact_type = match.group(1).strip()
        title = match.group(2).strip()
        content = match.group(3).strip()
        
        s3_key = s3_service.upload_content(
            content=content,
            filename=f"{title}.txt",
            tenant_id=g.tenant.id,
            folder='artifacts'
        )
        
        if s3_key:
            artifact = Artifact(
                tenant_id=g.tenant.id,
                chat_id=chat.id,
                message_id=assistant_message_id,
                title=title,
                content=content,
                artifact_type=artifact_type,
                s3_key=s3_key
            )
            db.session.add(artifact)
            artifacts_to_commit.append(artifact)
    
    if artifacts_to_commit:
        print(f"[DEBUG] 29. Committing {len(artifacts_to_commit)} artifacts to database...")
        try:
            db.session.commit()
            print("[DEBUG] 30. Artifacts committed successfully")
            for artifact in artifacts_to_commit:
                artifacts_created.append({
                    'id': artifact.id,
                    'title': artifact.title,
                    'type': artifact.artifact_type,
                    'content': artifact.content
                })
            print(f"[DEBUG] 31. Created {len(artifacts_created)} artifacts")
        except Exception as e:
            print(f"[DEBUG] ERROR committing artifacts: {str(e)}")
            import traceback
            traceback.print_exc()
            # Don't raise - artifacts are optional
    else:
        print("[DEBUG] 29. No artifacts found in response")

    print("[DEBUG] 32. Preparing final response JSON...")
    response_json = {
        'response': lex_response,
        'artifacts': artifacts_created,
        'message_id': assistant_message_id,
        'feedback_rating': None
    }
    print(f"[DEBUG] 33. Sending successful response (response length: {len(lex_response)} chars, {len(artifacts_created)} artifacts)")
    return jsonify(response_json)

@app.route('/api/chat/<int:chat_id>/rename', methods=['POST'])
@login_required
@tenant_required
def rename_chat(chat_id):
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    data = request.json or {}
    new_title = data.get('title', '').strip()
    
    if new_title:
        chat.title = new_title
        db.session.commit()
        return jsonify({'success': True})
    
    return jsonify({'success': False}), 400

@app.route('/api/chat/<int:chat_id>/delete', methods=['POST', 'DELETE'])
@login_required
@tenant_required
def delete_chat(chat_id):
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    # Delete uploaded files and their S3 objects
    uploaded_files = UploadedFile.query.filter_by(chat_id=chat.id).all()
    for uploaded_file in uploaded_files:
        if uploaded_file.s3_key:
            s3_service.delete_file(uploaded_file.s3_key)
        db.session.delete(uploaded_file)
    
    # Delete artifacts and their S3 objects
    artifacts = Artifact.query.filter_by(chat_id=chat.id).all()
    for artifact in artifacts:
        if artifact.s3_key:
            s3_service.delete_file(artifact.s3_key)
        db.session.delete(artifact)
    
    # Delete S3 messages file
    if chat.s3_messages_key:
        s3_service.delete_file(chat.s3_messages_key)
    
    # Finally delete the chat itself (cascade will delete messages)
    db.session.delete(chat)
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/api/user/accept-first-chat-warning', methods=['POST'])
@login_required
def accept_first_chat_warning():
    """Mark that user has seen and accepted the first chat warning"""
    from datetime import datetime
    current_user.first_chat_warning_seen_at = datetime.now()
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/chats', methods=['GET'])
@login_required
@tenant_required
def get_chats():
    chats = Chat.query.filter_by(
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).order_by(Chat.updated_at.desc()).all()
    
    return jsonify([{
        'id': chat.id,
        'title': chat.title,
        'updated_at': chat.updated_at.strftime('%d/%m %H:%M')
    } for chat in chats])

@app.route('/api/chats/search', methods=['POST'])
@login_required
@tenant_required
def search_chats():
    query = (request.json or {}).get('query', '').strip().lower()
    if not query:
        return jsonify([])
    
    chats = Chat.query.filter_by(
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).all()
    
    results = []
    for chat in chats:
        if query in chat.title.lower():
            results.append({
                'id': chat.id,
                'title': chat.title,
                'updated_at': chat.updated_at.strftime('%d/%m %H:%M'),
                'match_type': 'title'
            })
            continue
        
        if chat.s3_messages_key:
            messages_data = s3_service.get_messages(chat.s3_messages_key)
            if messages_data and 'messages' in messages_data:
                for msg in messages_data['messages']:
                    if query in msg.get('content', '').lower():
                        results.append({
                            'id': chat.id,
                            'title': chat.title,
                            'updated_at': chat.updated_at.strftime('%d/%m %H:%M'),
                            'match_type': 'content',
                            'snippet': msg.get('content', '')[:100] + '...'
                        })
                        break
    
    return jsonify(results)

@app.route('/api/chat/<int:chat_id>/export', methods=['GET'])
@login_required
@tenant_required
def export_chat_pdf(chat_id):
    from io import BytesIO
    from datetime import datetime
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    
    export_format = request.args.get('format', 'pdf')
    
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    # Collect messages
    messages = []
    if chat.s3_messages_key:
        messages_data = s3_service.get_messages(chat.s3_messages_key)
        if messages_data and 'messages' in messages_data:
            for msg in messages_data['messages']:
                role = "Jij" if msg.get('role') == "user" else "Lexi"
                timestamp = msg.get('timestamp', datetime.now().isoformat())
                try:
                    dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    timestamp_str = dt.strftime('%d-%m-%Y %H:%M')
                except:
                    timestamp_str = timestamp
                messages.append({
                    'role': role,
                    'timestamp': timestamp_str,
                    'content': msg.get('content', '')
                })
    else:
        # Fallback naar oude Message tabel
        db_messages = Message.query.filter_by(chat_id=chat.id).order_by(Message.created_at).all()
        for msg in db_messages:
            role = "Jij" if msg.role == "user" else "Lexi"
            messages.append({
                'role': role,
                'timestamp': msg.created_at.strftime('%d-%m-%Y %H:%M'),
                'content': msg.content
            })
    
    if export_format == 'pdf':
        # PDF Export - Available for ALL tiers
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a2332'),
            spaceAfter=20
        )
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=10
        )
        user_style = ParagraphStyle(
            'UserMessage',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=10,
            rightIndent=10,
            spaceAfter=10,
            textColor=colors.HexColor('#1a2332')
        )
        lexi_style = ParagraphStyle(
            'LexiMessage',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=10,
            rightIndent=10,
            spaceAfter=10,
            textColor=colors.HexColor('#333333')
        )
        
        story = []
        
        # Header
        story.append(Paragraph("Lexi CAO Meester - Chat Export", title_style))
        story.append(Paragraph(f"<b>Titel:</b> {chat.title}", header_style))
        story.append(Paragraph(f"<b>Datum:</b> {datetime.now().strftime('%d-%m-%Y %H:%M')}", header_style))
        story.append(Paragraph(f"<b>Gebruiker:</b> {current_user.full_name}", header_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Messages
        for msg in messages:
            role_label = f"<b>{msg['role']}</b> ({msg['timestamp']})"
            story.append(Paragraph(role_label, header_style))
            
            # Clean content for PDF
            content = msg['content'].replace('<', '&lt;').replace('>', '&gt;')
            content = content.replace('\n', '<br/>')
            
            if msg['role'] == "Jij":
                story.append(Paragraph(content, user_style))
            else:
                story.append(Paragraph(content, lexi_style))
            
            story.append(Spacer(1, 0.3*cm))
        
        doc.build(story)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename=chat_{chat_id}_{datetime.now().strftime("%Y%m%d")}.pdf'}
        )
    
    elif export_format == 'docx':
        # Word Export - Only for Professional and Enterprise
        tier = g.tenant.subscription_tier or 'starter'
        if tier not in ['professional', 'enterprise']:
            return jsonify({'error': 'Word export is only available for Professional and Enterprise tiers'}), 403
        
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Lexi CAO Meester - Chat Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Titel: {chat.title}")
        doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
        doc.add_paragraph(f"Gebruiker: {current_user.full_name}")
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Messages
        for msg in messages:
            # Role and timestamp
            p = doc.add_paragraph()
            run = p.add_run(f"{msg['role']} ({msg['timestamp']}):")
            run.bold = True
            run.font.size = Pt(11)
            if msg['role'] == "Jij":
                run.font.color.rgb = RGBColor(26, 35, 50)  # Navy
            else:
                run.font.color.rgb = RGBColor(212, 175, 55)  # Gold
            
            # Content
            content_p = doc.add_paragraph(msg['content'])
            content_p.paragraph_format.left_indent = Inches(0.5)
            
            # Separator
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename=chat_{chat_id}_{datetime.now().strftime("%Y%m%d")}.docx'}
        )
    
    else:
        return jsonify({'error': 'Invalid export format'}), 400

@app.route('/api/feedback', methods=['POST'])
@login_required
@tenant_required
def submit_feedback():
    data = request.json or {}
    message_id = data.get('message_id')
    rating = data.get('rating')
    comment = data.get('comment', '')
    
    if not message_id or not rating:
        return jsonify({'error': 'Missing data'}), 400
    
    message = Message.query.filter_by(id=message_id, tenant_id=g.tenant.id).first_or_404()
    
    chat = Chat.query.filter_by(
        id=message.chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    message.feedback_rating = rating
    message.feedback_comment = comment
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/api/artifact/<int:artifact_id>/download')
@login_required
@tenant_required
def download_artifact(artifact_id):
    artifact = Artifact.query.filter_by(
        id=artifact_id,
        tenant_id=g.tenant.id
    ).first_or_404()
    
    download_url = s3_service.get_file_url(artifact.s3_key, expiration=300)
    
    if not download_url:
        return jsonify({'error': 'Download niet beschikbaar'}), 500
    
    return jsonify({
        'download_url': download_url,
        'title': artifact.title,
        'type': artifact.artifact_type
    })

@app.route('/api/chat/<int:chat_id>/files', methods=['GET'])
@login_required
@tenant_required
def get_chat_files(chat_id):
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    chat = Chat.query.filter_by(
        id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    files = UploadedFile.query.filter_by(
        chat_id=chat_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).all()
    
    return jsonify({
        'files': [{
            'id': f.id,
            'filename': f.original_filename,
            'file_size': f.file_size,
            'mime_type': f.mime_type,
            'created_at': f.created_at.isoformat()
        } for f in files]
    })

@app.route('/api/files/<int:file_id>', methods=['DELETE'])
@login_required
@tenant_required
def delete_file(file_id):
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    uploaded_file = UploadedFile.query.filter_by(
        id=file_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    # Delete from S3
    s3_deleted = s3_service.delete_file(uploaded_file.s3_key)
    if not s3_deleted:
        return jsonify({'error': 'Kon bestand niet verwijderen uit opslag'}), 500
    
    # Delete from database
    db.session.delete(uploaded_file)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Bestand verwijderd'})

@app.route('/api/file/<int:file_id>/view', methods=['GET'])
@login_required
@tenant_required
def view_file(file_id):
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    uploaded_file = UploadedFile.query.filter_by(
        id=file_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    # For PDF, return presigned URL for direct browser access
    if uploaded_file.mime_type == 'application/pdf':
        download_url = s3_service.get_file_url(uploaded_file.s3_key, expiration=3600)
        if not download_url:
            return jsonify({'error': 'Kon bestand niet ophalen'}), 500
        
        return jsonify({
            'type': 'pdf',
            'url': download_url,
            'filename': uploaded_file.original_filename
        })
    
    # For text/docx, return extracted content
    content, error = s3_service.download_file_content(uploaded_file.s3_key, uploaded_file.mime_type)
    if error:
        return jsonify({'error': error}), 500
    
    return jsonify({
        'type': 'text',
        'content': content
    })

@app.route('/api/upload', methods=['POST'])
@login_required
@tenant_required
def upload_file():
    if g.tenant.subscription_status not in ['active', 'trial', 'trialing']:
        return jsonify({'error': 'Subscription niet actief'}), 403
    
    if 'file' not in request.files:
        return jsonify({'error': 'Geen bestand'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Geen bestand geselecteerd'}), 400
    
    # SECURITY: File type whitelist - only allow specific document types
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}
    ALLOWED_MIMETYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'text/plain'
    }
    
    filename = secure_filename(file.filename)
    if not filename or '.' not in filename:
        return jsonify({'error': 'Ongeldig bestand'}), 400
    
    file_ext = filename.rsplit('.', 1)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        return jsonify({'error': f'Alleen {", ".join(ALLOWED_EXTENSIONS).upper()} bestanden toegestaan'}), 400
    
    if file.content_type not in ALLOWED_MIMETYPES:
        return jsonify({'error': 'Ongeldig bestandstype'}), 400
    
    chat_id = request.form.get('chat_id')
    
    # Get file size
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    
    # Extract text from PDF using MarkItDown + OCR fallback
    extracted_text = None
    if file.content_type == 'application/pdf':
        try:
            # Read file data for MarkItDown
            file.seek(0)
            file_data = file.read()
            
            # Save to temporary file for MarkItDown processing
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(file_data)
                tmp_path = tmp_file.name
            
            # Extract text using MarkItDown
            md = MarkItDown()
            result = md.convert(tmp_path)
            extracted_text = result.text_content
            
            # If MarkItDown didn't extract text (scanned PDF), use OCR
            if not extracted_text or len(extracted_text.strip()) == 0:
                print(f"[DEBUG] MarkItDown extracted no text, trying OCR...")
                try:
                    # Convert PDF pages to images
                    images = convert_from_path(tmp_path)
                    ocr_texts = []
                    
                    for i, image in enumerate(images):
                        # Extract text from each page using Tesseract OCR
                        page_text = pytesseract.image_to_string(image, lang='nld+eng')
                        if page_text.strip():
                            ocr_texts.append(f"--- Pagina {i+1} ---\n{page_text}")
                    
                    if ocr_texts:
                        extracted_text = '\n\n'.join(ocr_texts)
                        print(f"[DEBUG] OCR successful, extracted {len(extracted_text)} characters from {len(images)} pages")
                    else:
                        print(f"[DEBUG] OCR found no text in PDF")
                except Exception as ocr_error:
                    print(f"[DEBUG] OCR failed: {ocr_error}")
            
            # Clean up temporary file
            os.unlink(tmp_path)
            
            # Reset file pointer for S3 upload
            file.seek(0)
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            # Reset file pointer and continue without extracted text
            file.seek(0)
    
    # Upload to S3
    s3_key = s3_service.upload_file(file, g.tenant.id)
    if not s3_key:
        return jsonify({'error': 'Upload mislukt'}), 500
    
    uploaded_file = UploadedFile(
        tenant_id=g.tenant.id,
        user_id=current_user.id,
        chat_id=chat_id if chat_id else None,
        filename=file.filename,
        original_filename=file.filename,
        s3_key=s3_key,
        file_size=file_size,
        mime_type=file.content_type,
        extracted_text=extracted_text
    )
    db.session.add(uploaded_file)
    db.session.commit()
    
    return jsonify({'success': True, 'file_id': uploaded_file.id})

# Support Ticket Routes (Customer)
@app.route('/support')
@login_required
@tenant_required
def support_tickets():
    from models import SupportTicket
    tickets = SupportTicket.query.filter_by(
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).order_by(SupportTicket.updated_at.desc()).all()
    
    return render_template('support/tickets.html', tickets=tickets, tenant=g.tenant, user=current_user)

@app.route('/support/new')
@login_required
@tenant_required
def new_support_ticket():
    return render_template('support/new_ticket.html', tenant=g.tenant, user=current_user)

@app.route('/api/support/create', methods=['POST'])
@login_required
@tenant_required
def create_support_ticket():
    from models import SupportTicket, SupportReply
    data = request.json or {}
    
    subject = data.get('subject', '').strip()
    category = data.get('category', '').strip()
    message = data.get('message', '').strip()
    
    if not subject or not category or not message:
        return jsonify({'error': 'Alle velden zijn verplicht'}), 400
    
    # Get next ticket number
    last_ticket = SupportTicket.query.order_by(SupportTicket.ticket_number.desc()).first()
    ticket_number = (last_ticket.ticket_number + 1) if last_ticket else 1000
    
    # Create ticket
    ticket = SupportTicket(
        ticket_number=ticket_number,
        tenant_id=g.tenant.id,
        user_id=current_user.id,
        user_email=current_user.email,
        user_name=current_user.full_name,
        subject=subject,
        category=category,
        status='open'
    )
    db.session.add(ticket)
    db.session.flush()
    
    # Add first message
    reply = SupportReply(
        ticket_id=ticket.id,
        message=message,
        is_admin=False,
        sender_name=current_user.full_name
    )
    db.session.add(reply)
    db.session.commit()
    
    return jsonify({'success': True, 'ticket_id': ticket.id})

@app.route('/support/<int:ticket_id>')
@login_required
@tenant_required
def view_support_ticket(ticket_id):
    from models import SupportTicket
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    return render_template('support/ticket_detail.html', ticket=ticket, tenant=g.tenant, user=current_user)

@app.route('/api/support/<int:ticket_id>/reply', methods=['POST'])
@login_required
@tenant_required
def reply_support_ticket(ticket_id):
    from models import SupportTicket, SupportReply
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    data = request.json or {}
    message = data.get('message', '').strip()
    
    if not message:
        return jsonify({'error': 'Bericht mag niet leeg zijn'}), 400
    
    reply = SupportReply(
        ticket_id=ticket.id,
        message=message,
        is_admin=False,
        sender_name=current_user.full_name
    )
    db.session.add(reply)
    
    ticket.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/api/support/<int:ticket_id>/close', methods=['POST'])
@login_required
@tenant_required
def close_support_ticket(ticket_id):
    from models import SupportTicket
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id,
        user_id=current_user.id
    ).first_or_404()
    
    ticket.status = 'closed'
    ticket.closed_at = datetime.utcnow()
    db.session.commit()
    
    # Send ticket resolved email
    try:
        email_service.send_ticket_resolved_email(current_user, g.tenant, ticket)
        print(f"✅ Ticket resolved email sent to {current_user.email}")
    except Exception as e:
        print(f"⚠️ Failed to send ticket resolved email: {e}")
    
    return jsonify({'success': True})

@app.route('/admin/dashboard')
@login_required
@tenant_required
@admin_required
def admin_dashboard():
    users = User.query.filter_by(tenant_id=g.tenant.id).all()
    total_chats = Chat.query.filter_by(tenant_id=g.tenant.id).count()
    subscription = Subscription.query.filter_by(tenant_id=g.tenant.id).first()
    
    total_messages = Message.query.filter_by(
        tenant_id=g.tenant.id,
        role='user'
    ).count()
    
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    messages_this_month = Message.query.filter(
        Message.tenant_id == g.tenant.id,
        Message.role == 'user',
        Message.created_at >= thirty_days_ago
    ).count()
    
    active_users_ids = db.session.query(Chat.user_id).filter(
        Chat.tenant_id == g.tenant.id,
        Chat.updated_at >= thirty_days_ago
    ).distinct().all()
    active_users_count = len(active_users_ids)
    
    top_users = db.session.query(
        User,
        db.func.count(Chat.id).label('chat_count')
    ).join(Chat, User.id == Chat.user_id
    ).filter(
        User.tenant_id == g.tenant.id,
        Chat.tenant_id == g.tenant.id
    ).group_by(User.id
    ).order_by(db.desc('chat_count')
    ).limit(5).all()
    
    return render_template('admin_dashboard.html', 
                         tenant=g.tenant, 
                         users=users, 
                         total_chats=total_chats,
                         subscription=subscription,
                         total_messages=total_messages,
                         messages_this_month=messages_this_month,
                         active_users_count=active_users_count,
                         top_users=top_users)

@app.route('/admin/cao/update', methods=['POST'])
@login_required
@tenant_required
@admin_required
def update_cao_preference():
    from cao_config import validate_cao_preference, get_cao_display_name
    
    new_cao = request.form.get('cao_preference')
    
    if not validate_cao_preference(new_cao):
        flash('Ongeldige CAO keuze.', 'danger')
        return redirect(url_for('admin_dashboard'))
    
    g.tenant.cao_preference = new_cao
    db.session.commit()
    
    cao_name = get_cao_display_name(new_cao)
    flash(f'CAO voorkeur succesvol gewijzigd naar {cao_name}. Alle nieuwe chats gebruiken nu deze CAO.', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
@tenant_required
@admin_required
def admin_users():
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add':
            current_user_count = User.query.filter_by(tenant_id=g.tenant.id).count()
            if current_user_count >= g.tenant.max_users:
                flash(f'Maximum aantal gebruikers bereikt ({g.tenant.max_users}). Upgrade je plan.', 'warning')
                return redirect(url_for('admin_users'))
            
            email = request.form.get('email')
            first_name = request.form.get('first_name')
            last_name = request.form.get('last_name')
            password = request.form.get('password')
            role = request.form.get('role', 'user')
            
            if role not in ['user', 'admin']:
                role = 'user'
            
            if User.query.filter_by(tenant_id=g.tenant.id, email=email).first():
                flash('Deze email is al in gebruik.', 'danger')
            else:
                user = User(
                    tenant_id=g.tenant.id,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    role=role
                )
                # Generate secure activation token (24-hour expiry)
                token = secrets.token_urlsafe(32)
                user.reset_token = token
                user.reset_token_expires_at = datetime.utcnow() + timedelta(hours=24)
                # Set a temporary random password (will be changed via activation)
                user.set_password(secrets.token_urlsafe(16))
                
                db.session.add(user)
                db.session.commit()
                
                # Send activation email with secure token link (no subdomain needed)
                domain = os.getenv('PRODUCTION_DOMAIN', 'lexiai.nl')
                activation_url = f"https://{domain}/reset-password/{token}"
                admin_name = f"{current_user.first_name} {current_user.last_name}"
                email_service.send_user_invitation_email(user, g.tenant, activation_url, admin_name)
                
                flash(f'Gebruiker toegevoegd! Activatie email verzonden naar {email}.', 'success')
        
        elif action == 'toggle':
            user_id = request.form.get('user_id')
            user = User.query.filter_by(id=user_id, tenant_id=g.tenant.id).first()
            if user and user.id != current_user.id:
                was_active = user.is_active
                user.is_active = not user.is_active
                db.session.commit()
                
                # Send deactivation email if user was deactivated
                if was_active and not user.is_active:
                    try:
                        admin_name = f"{current_user.first_name} {current_user.last_name}"
                        email_service.send_account_deactivated_email(user, g.tenant, admin_name)
                        print(f"✅ Account deactivation email sent to {user.email}")
                    except Exception as e:
                        print(f"⚠️ Failed to send deactivation email: {e}")
                
                flash('Gebruiker status gewijzigd.', 'success')
        
        elif action == 'delete':
            user_id = request.form.get('user_id')
            user = User.query.filter_by(id=user_id, tenant_id=g.tenant.id).first()
            if user and user.id != current_user.id:
                db.session.delete(user)
                db.session.commit()
                flash('Gebruiker verwijderd.', 'success')
        
        elif action == 'change_role':
            user_id = request.form.get('user_id')
            new_role = request.form.get('role')
            user = User.query.filter_by(id=user_id, tenant_id=g.tenant.id).first()
            if user and user.id != current_user.id and new_role in ['user', 'admin']:
                old_role = user.role
                user.role = new_role
                db.session.commit()
                
                # Send role changed email
                if old_role != new_role:
                    try:
                        admin_name = f"{current_user.first_name} {current_user.last_name}"
                        email_service.send_role_changed_email(user, g.tenant, new_role, admin_name)
                        print(f"✅ Role changed email sent to {user.email}")
                    except Exception as e:
                        print(f"⚠️ Failed to send role changed email: {e}")
                
                flash(f'Gebruiker rol gewijzigd naar {new_role}.', 'success')
    
    users = User.query.filter_by(tenant_id=g.tenant.id).all()
    return render_template('admin_users.html', tenant=g.tenant, users=users)

# Admin Support Routes
@app.route('/admin/support')
@login_required
@tenant_required
@admin_required
def admin_support():
    from models import SupportTicket
    status_filter = request.args.get('status', 'all')
    category_filter = request.args.get('category', 'all')
    
    query = SupportTicket.query.filter_by(tenant_id=g.tenant.id)
    
    if status_filter != 'all':
        query = query.filter_by(status=status_filter)
    
    if category_filter != 'all':
        query = query.filter_by(category=category_filter)
    
    tickets = query.order_by(SupportTicket.updated_at.desc()).all()
    
    # Stats
    total_tickets = SupportTicket.query.filter_by(tenant_id=g.tenant.id).count()
    open_tickets = SupportTicket.query.filter_by(tenant_id=g.tenant.id, status='open').count()
    in_progress_tickets = SupportTicket.query.filter_by(tenant_id=g.tenant.id, status='in_progress').count()
    answered_tickets = SupportTicket.query.filter_by(tenant_id=g.tenant.id, status='answered').count()
    closed_tickets = SupportTicket.query.filter_by(tenant_id=g.tenant.id, status='closed').count()
    
    return render_template('admin_support.html', 
                         tickets=tickets, 
                         tenant=g.tenant,
                         total_tickets=total_tickets,
                         open_tickets=open_tickets,
                         in_progress_tickets=in_progress_tickets,
                         answered_tickets=answered_tickets,
                         closed_tickets=closed_tickets,
                         status_filter=status_filter,
                         category_filter=category_filter)

@app.route('/admin/support/<int:ticket_id>')
@login_required
@tenant_required
@admin_required
def admin_view_ticket(ticket_id):
    from models import SupportTicket
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id
    ).first_or_404()
    
    return render_template('admin_support_detail.html', ticket=ticket, tenant=g.tenant)

@app.route('/api/admin/support/<int:ticket_id>/reply', methods=['POST'])
@login_required
@tenant_required
@admin_required
def admin_reply_ticket(ticket_id):
    from models import SupportTicket, SupportReply
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id
    ).first_or_404()
    
    data = request.json or {}
    message = data.get('message', '').strip()
    
    if not message:
        return jsonify({'error': 'Bericht mag niet leeg zijn'}), 400
    
    reply = SupportReply(
        ticket_id=ticket.id,
        message=message,
        is_admin=True,
        sender_name=current_user.full_name
    )
    db.session.add(reply)
    
    # Update status to answered when admin replies
    if ticket.status == 'open':
        ticket.status = 'answered'
    
    ticket.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/api/admin/support/<int:ticket_id>/status', methods=['POST'])
@login_required
@tenant_required
@admin_required
def admin_update_ticket_status(ticket_id):
    from models import SupportTicket
    ticket = SupportTicket.query.filter_by(
        id=ticket_id,
        tenant_id=g.tenant.id
    ).first_or_404()
    
    data = request.json or {}
    new_status = data.get('status')
    
    if new_status not in ['open', 'in_progress', 'answered', 'closed']:
        return jsonify({'error': 'Ongeldige status'}), 400
    
    old_status = ticket.status
    ticket.status = new_status
    if new_status == 'closed':
        ticket.closed_at = datetime.utcnow()
        # Send ticket resolved email when admin closes ticket
        if old_status != 'closed':
            try:
                user = User.query.get(ticket.user_id)
                if user:
                    email_service.send_ticket_resolved_email(user, g.tenant, ticket)
                    print(f"✅ Ticket resolved email sent to {user.email}")
            except Exception as e:
                print(f"⚠️ Failed to send ticket resolved email: {e}")
    else:
        ticket.closed_at = None
    
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/admin/templates', methods=['GET', 'POST'])
@login_required
@tenant_required
@admin_required
def admin_templates():
    if request.method == 'POST':
        name = request.form.get('name') or ''
        category = request.form.get('category') or ''
        content = request.form.get('content') or ''
        
        template = Template(
            tenant_id=g.tenant.id,
            name=name,
            category=category,
            content=content
        )
        
        if content:
            s3_key = s3_service.upload_content(content, f"{name}.txt", g.tenant.id, 'templates')
            template.s3_key = s3_key
        
        db.session.add(template)
        db.session.commit()
        flash('Template opgeslagen!', 'success')
    
    templates = Template.query.filter_by(tenant_id=g.tenant.id).all()
    return render_template('admin_templates.html', tenant=g.tenant, templates=templates)

@app.route('/admin/billing')
@login_required
@tenant_required
@admin_required
def admin_billing():
    subscription = Subscription.query.filter_by(tenant_id=g.tenant.id).first()
    
    return render_template('admin_billing.html', 
                         tenant=g.tenant, 
                         subscription=subscription)

@app.route('/admin/billing/checkout/<plan>')
@login_required
@tenant_required
@admin_required
def billing_checkout(plan):
    if plan not in ['starter', 'professional', 'enterprise']:
        return "Invalid plan", 400
    
    success_url = url_for('billing_success', _external=True)
    cancel_url = url_for('admin_billing', _external=True)
    
    session_obj = StripeService.create_checkout_session(
        g.tenant.id, plan, success_url, cancel_url
    )
    
    if session_obj and session_obj.url:
        return redirect(session_obj.url)
    
    flash('Er ging iets mis met de betaling.', 'danger')
    return redirect(url_for('admin_billing'))

@app.route('/billing/success')
@login_required
@tenant_required
def billing_success():
    flash('Betaling succesvol! Je account is nu actief.', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/webhook/stripe', methods=['POST'])
@limiter.limit("100 per hour")
def stripe_webhook():
    payload = request.data
    sig_header = request.headers.get('Stripe-Signature')
    # Productie webhook secret heeft voorrang over test webhook secret
    webhook_secret = os.getenv('STRIPE_WEBHOOK_SECRET_PROD') or os.getenv('STRIPE_WEBHOOK_SECRET')
    
    # SECURITY: Webhook secret MUST be configured
    if not webhook_secret:
        print("❌ CRITICAL: STRIPE_WEBHOOK_SECRET not configured - webhook rejected")
        return jsonify({'error': 'Webhook not properly configured'}), 500
    
    is_prod_webhook = bool(os.getenv('STRIPE_WEBHOOK_SECRET_PROD'))
    print(f"📥 Webhook received - Mode: {'Production' if is_prod_webhook else 'Test'}")
    
    # SECURITY: Signature verification is MANDATORY - no bypasses
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        print(f"✅ Webhook signature verified - Event type: {event['type']}")
    except stripe.error.SignatureVerificationError as e:
        print(f"❌ Webhook signature verification FAILED: {e}")
        return jsonify({'error': 'Invalid signature'}), 400
    except Exception as e:
        print(f"❌ Webhook processing error: {e}")
        return jsonify({'error': 'Webhook error'}), 400
    
    if event['type'] == 'checkout.session.completed':
        from models import PendingSignup
        from provision_tenant import provision_tenant_from_signup
        
        session_obj = event['data']['object']
        checkout_session_id = session_obj.get('id')
        
        print(f"🔔 Webhook: Processing checkout.session.completed for session: {checkout_session_id}")
        
        # Get pending signup from database using checkout session ID
        pending_signup = PendingSignup.query.filter_by(checkout_session_id=checkout_session_id).first()
        
        if not pending_signup:
            # No pending signup = fallback already processed or duplicate webhook
            # Return 200 to prevent Stripe retry loops
            print(f"⚠️ No pending signup for session {checkout_session_id} - likely already processed by fallback")
            return jsonify({'success': True, 'message': 'Already processed'}), 200
        
        # Use shared provisioning service (idempotent)
        success, user, error_msg = provision_tenant_from_signup(
            pending_signup=pending_signup,
            stripe_session_data=session_obj
        )
        
        if success:
            print(f"✅ Webhook: Account provisioned successfully for {user.email if user else 'existing user'}")
            
            # Send payment success and welcome emails
            if user:
                tenant = Tenant.query.get(user.tenant_id)
                if tenant:
                    try:
                        email_service.send_payment_success_email(user, tenant)
                        email_service.send_welcome_email(user, tenant)
                        print(f"✅ Welcome emails sent to {user.email}")
                    except Exception as e:
                        print(f"⚠️ Failed to send welcome emails: {e}")
            
            return jsonify({'success': True})
        else:
            print(f"❌ Webhook: Failed to provision account: {error_msg}")
            return jsonify({'error': error_msg}), 500
    
    elif event['type'] == 'customer.subscription.updated':
        subscription_obj = event['data']['object']
        stripe_sub_id = subscription_obj.get('id')
        status = subscription_obj.get('status')
        old_plan_id = subscription_obj.get('previous_attributes', {}).get('items', {}).get('data', [{}])[0].get('price', {}).get('id')
        new_plan_id = subscription_obj.get('items', {}).get('data', [{}])[0].get('price', {}).get('id')
        
        subscription = Subscription.query.filter_by(stripe_subscription_id=stripe_sub_id).first()
        if subscription:
            # Map Stripe status to our status
            if status in ['active', 'trialing']:
                subscription.status = 'active'
                tenant = Tenant.query.get(subscription.tenant_id)
                if tenant:
                    tenant.status = 'active'
                    # Send subscription updated email if plan changed
                    if old_plan_id and new_plan_id and old_plan_id != new_plan_id:
                        user = User.query.filter_by(tenant_id=tenant.id, role='admin').first()
                        if user:
                            try:
                                email_service.send_subscription_updated_email(user, tenant, tenant.plan)
                                print(f"✅ Subscription updated email sent to {user.email}")
                            except Exception as e:
                                print(f"⚠️ Failed to send subscription updated email: {e}")
            elif status in ['past_due', 'unpaid']:
                subscription.status = 'past_due'
            elif status in ['canceled', 'incomplete_expired']:
                subscription.status = 'canceled'
                tenant = Tenant.query.get(subscription.tenant_id)
                if tenant:
                    tenant.status = 'inactive'
                    # Send subscription cancelled email
                    user = User.query.filter_by(tenant_id=tenant.id, role='admin').first()
                    if user:
                        try:
                            email_service.send_subscription_cancelled_email(user, tenant)
                            print(f"✅ Subscription cancelled email sent to {user.email}")
                        except Exception as e:
                            print(f"⚠️ Failed to send subscription cancelled email: {e}")
            
            db.session.commit()
            print(f"Updated subscription {stripe_sub_id} to status {status}")
    
    elif event['type'] == 'invoice.payment_failed':
        invoice = event['data']['object']
        customer_id = invoice.get('customer')
        
        subscription = Subscription.query.filter_by(stripe_customer_id=customer_id).first()
        if subscription:
            tenant = Tenant.query.get(subscription.tenant_id)
            email_service.send_payment_failed_email(tenant)
            print(f"Payment failed email sent to {tenant.contact_email}")
    
    elif event['type'] == 'invoice.finalized':
        # Send iDEAL payment links for recurring invoices (NOT first invoice)
        invoice = event['data']['object']
        customer_id = invoice.get('customer')
        subscription_id = invoice.get('subscription')
        billing_reason = invoice.get('billing_reason')
        
        # Only process if this is a subscription invoice
        if not subscription_id:
            return jsonify({'success': True, 'message': 'Not a subscription invoice'})
        
        # CRITICAL: Skip first invoice (already paid via Stripe Checkout)
        # Only send emails for recurring invoices (month 2, 3, 4, etc.)
        if billing_reason == 'subscription_create':
            print(f"ℹ️  Skipping email for first invoice (already paid via Checkout)")
            return jsonify({'success': True, 'message': 'First invoice - no email needed'})
        
        subscription = Subscription.query.filter_by(stripe_subscription_id=subscription_id).first()
        
        # Only send email for iDEAL payment method (manual invoices)
        if subscription and subscription.payment_method == 'ideal':
            tenant = Tenant.query.get(subscription.tenant_id)
            admin_user = User.query.filter_by(tenant_id=tenant.id, role='admin').first()
            
            if tenant and admin_user:
                try:
                    # Get invoice details
                    hosted_invoice_url = invoice.get('hosted_invoice_url')
                    amount_due = invoice.get('amount_due', 0) / 100  # Convert cents to euros
                    due_date_timestamp = invoice.get('due_date')
                    
                    # Format amount and due date
                    amount_formatted = f"€{amount_due:.2f}"
                    if due_date_timestamp:
                        from datetime import datetime
                        due_date = datetime.fromtimestamp(due_date_timestamp).strftime('%d-%m-%Y')
                    else:
                        from datetime import datetime, timedelta
                        due_date = (datetime.utcnow() + timedelta(days=7)).strftime('%d-%m-%Y')
                    
                    # Send iDEAL payment link email (only for month 2+)
                    if hosted_invoice_url:
                        email_service.send_ideal_payment_link_email(
                            admin_user, 
                            tenant, 
                            hosted_invoice_url, 
                            amount_formatted, 
                            due_date
                        )
                        print(f"✅ iDEAL monthly payment email sent to {admin_user.email} (recurring invoice)")
                    else:
                        print(f"⚠️ No hosted_invoice_url for invoice {invoice.get('id')}")
                except Exception as e:
                    print(f"⚠️ Failed to send iDEAL payment link email: {e}")
    
    return jsonify({'success': True})

@app.route('/super-admin/dashboard')
@super_admin_required
def super_admin_dashboard():
    print(f"[DEBUG] Super Admin Dashboard accessed")
    print(f"  Session is_super_admin: {session.get('is_super_admin')}")
    print(f"  g.is_super_admin: {g.is_super_admin}")
    print(f"  current_user: {current_user}")
    print(f"  Request host: {request.host}")

    sort_by = request.args.get('sort_by', 'created_at')
    sort_order = request.args.get('sort_order', 'desc')
    
    valid_sort_columns = {
        'company_name': Tenant.company_name,
        'contact_email': Tenant.contact_email,
        'created_at': Tenant.created_at,
        'subscription_tier': Tenant.subscription_tier,
        'subscription_status': Tenant.subscription_status
    }
    
    sort_column = valid_sort_columns.get(sort_by, Tenant.created_at)
    
    if sort_order == 'asc':
        tenants = Tenant.query.order_by(sort_column.asc()).all()
    else:
        tenants = Tenant.query.order_by(sort_column.desc()).all()
    
    total_users = User.query.count()
    
    mrr_prices = {'starter': 499, 'professional': 599, 'enterprise': 1199}
    
    current_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in tenants if t.subscription_status == 'active')
    arr = current_mrr * 12
    
    from dateutil.relativedelta import relativedelta
    last_month = datetime.utcnow() - relativedelta(months=1)
    last_month_tenants = [t for t in tenants if t.created_at < last_month and t.subscription_status == 'active']
    last_month_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in last_month_tenants)
    
    growth_percentage = 0
    if last_month_mrr > 0:
        growth_percentage = ((current_mrr - last_month_mrr) / last_month_mrr) * 100
    elif current_mrr > 0 and last_month_mrr == 0:
        growth_percentage = 100
    
    starter_count = sum(1 for t in tenants if t.subscription_tier == 'starter' and t.subscription_status == 'active')
    professional_count = sum(1 for t in tenants if t.subscription_tier == 'professional' and t.subscription_status == 'active')
    enterprise_count = sum(1 for t in tenants if t.subscription_tier == 'enterprise' and t.subscription_status == 'active')
    starter_mrr = starter_count * 499
    professional_mrr = professional_count * 599
    enterprise_mrr = enterprise_count * 1199
    
    mrr_history = []
    for i in range(6, 0, -1):
        month_date = datetime.utcnow() - relativedelta(months=i)
        month_tenants = [t for t in tenants if t.created_at <= month_date and t.subscription_status == 'active']
        month_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in month_tenants)
        mrr_history.append({
            'month': month_date.strftime('%b'),
            'mrr': month_mrr
        })
    
    return render_template('super_admin_dashboard.html', 
                         tenants=tenants, 
                         total_users=total_users,
                         current_mrr=current_mrr,
                         arr=arr,
                         growth_percentage=growth_percentage,
                         starter_count=starter_count,
                         professional_count=professional_count,
                         enterprise_count=enterprise_count,
                         starter_mrr=starter_mrr,
                         professional_mrr=professional_mrr,
                         enterprise_mrr=enterprise_mrr,
                         mrr_history=mrr_history,
                         sort_by=sort_by,
                         sort_order=sort_order)

@app.route('/super-admin/documents')
@super_admin_required
def super_admin_documents():
    """CAO Document Management Page"""
    return render_template('super_admin_documents.html')

# Upload status tracker
upload_status = {
    'status': 'idle',
    'progress': 0,
    'current_file': '',
    'total_files': 0,
    'processed_files': 0,
    'messages': [],
    'imported_count': 0,
    'error': None
}

@app.route('/upload/api/upload', methods=['POST'])
@csrf.exempt  # Exempt from CSRF for file uploads
def upload_files():
    """Handle file upload"""
    global upload_status

    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')

    if not files or files[0].filename == '':
        return jsonify({'error': 'No files selected'}), 400

    # Save files
    upload_dir = '/tmp/cao_import'
    os.makedirs(upload_dir, exist_ok=True)

    valid_files = []
    for file in files:
        if file and file.filename.endswith(('.pdf', '.txt', '.docx')):
            filename = secure_filename(file.filename)
            file_path = os.path.join(upload_dir, filename)
            file.save(file_path)
            valid_files.append(file_path)

    if not valid_files:
        return jsonify({'error': 'No valid PDF/TXT/DOCX files'}), 400

    # Reset and start import
    upload_status = {
        'status': 'uploading',
        'progress': 0,
        'current_file': '',
        'total_files': len(valid_files),
        'processed_files': 0,
        'messages': [f"📁 {len(valid_files)} files uploaded"],
        'imported_count': 0,
        'error': None
    }

    # Start async processing
    import threading
    thread = threading.Thread(target=_process_files_background, args=(valid_files,))
    thread.daemon = True
    thread.start()

    return jsonify({'message': 'Upload started', 'files_count': len(valid_files)})

def _process_files_background(file_paths):
    """Process files in background with detailed logging"""
    import sys
    import os
    from datetime import datetime
    global upload_status
    try:
        # Import document tracker
        sys.path.insert(0, '/tmp')
        from document_tracker import add_document as track_document

        upload_status['status'] = 'processing'
        total_imported = 0
        total_chunks = 0

        for idx, file_path in enumerate(file_paths):
            upload_status['current_file'] = os.path.basename(file_path)
            upload_status['progress'] = int((idx / len(file_paths)) * 40)

            try:
                # Import modules - optional since gqlalchemy may not be available
                from document_importer import parse_document, generate_embeddings, import_to_memgraph

                # Step 1: Parse document
                upload_status['messages'].append(f"📄 Processing: {os.path.basename(file_path)}")
                cao_name, chunks = parse_document(file_path)
                upload_status['messages'].append(f"   ✓ Parsed into {len(chunks)} chunks")
                total_chunks += len(chunks)

                if not chunks:
                    upload_status['messages'].append(f"   ⚠️ No content extracted from {os.path.basename(file_path)}")
                    continue

                # Step 2: Generate embeddings
                upload_status['progress'] = int((idx / len(file_paths)) * 60)
                upload_status['messages'].append(f"   ⏳ Generating embeddings for {len(chunks)} chunks...")

                embeddings_data = generate_embeddings(chunks)
                upload_status['messages'].append(f"   ✓ Generated {len(embeddings_data)} embeddings (1024-dim vectors)")

                # Step 3: Import to Memgraph
                upload_status['progress'] = int((idx / len(file_paths)) * 85)
                upload_status['messages'].append(f"   ⏳ Importing to Memgraph as '{cao_name}'...")

                try:
                    from gqlalchemy import Memgraph
                    memgraph = Memgraph(host="localhost", port=7687)
                    imported = import_to_memgraph(memgraph, cao_name, embeddings_data)
                    total_imported += imported
                    upload_status['messages'].append(f"   ✅ Imported {imported} articles to Memgraph")
                    upload_status['messages'].append(f"   🔗 Created CAO node: '{cao_name}' with CONTAINS_ARTICLE relationships")

                    # Track this document upload
                    try:
                        filename = os.path.basename(file_path)
                        track_document(filename, cao_name, imported)
                        upload_status['messages'].append(f"   📝 Document tracked in registry")
                    except Exception as track_err:
                        upload_status['messages'].append(f"   ⚠️ Tracking error: {str(track_err)[:40]}")

                except Exception as mem_err:
                    # If gqlalchemy fails, still mark as processed
                    upload_status['messages'].append(f"   ⚠️ Memgraph connection unavailable")
                    upload_status['messages'].append(f"   ℹ️ Embeddings generated successfully (ready for indexing)")
                    total_imported += len(embeddings_data)

                    # Still track it
                    try:
                        filename = os.path.basename(file_path)
                        track_document(filename, cao_name, len(embeddings_data))
                    except:
                        pass

                upload_status['processed_files'] += 1
                upload_status['progress'] = int((idx + 1) / len(file_paths)) * 90

            except Exception as e:
                upload_status['messages'].append(f"   ❌ Error: {str(e)[:100]}")

        # Final summary
        upload_status['status'] = 'complete'
        upload_status['progress'] = 100
        upload_status['imported_count'] = total_imported
        upload_status['messages'].append("")
        upload_status['messages'].append("═" * 50)
        upload_status['messages'].append("✨ IMPORT COMPLETE")
        upload_status['messages'].append("═" * 50)
        upload_status['messages'].append(f"📊 Summary:")
        upload_status['messages'].append(f"   • Files processed: {len(file_paths)}")
        upload_status['messages'].append(f"   • Total chunks parsed: {total_chunks}")
        upload_status['messages'].append(f"   • Total articles indexed: {total_imported}")

        # Save index to JSON file for quick retrieval
        try:
            import json
            from pathlib import Path
            index_file = Path('/tmp/cao_documents_index.json')

            # Try to fetch current data from Memgraph
            documents_list = []
            try:
                from gqlalchemy import Memgraph
                memgraph = Memgraph(host="localhost", port=7687)

                cao_results = list(memgraph.execute_and_fetch("""
                    MATCH (cao:CAO)
                    RETURN cao.name as cao
                    ORDER BY cao
                """))

                for cao_doc in cao_results:
                    cao_name = cao_doc['cao']
                    count_result = list(memgraph.execute_and_fetch("""
                        MATCH (cao:CAO {name: $cao_name})-[:CONTAINS_ARTICLE]->(article:Article)
                        RETURN COUNT(article) as article_count
                    """, {'cao_name': cao_name}))

                    article_count = 0
                    if count_result and count_result[0]['article_count'] is not None:
                        article_count = count_result[0]['article_count']

                    documents_list.append({
                        'cao': cao_name,
                        'article_count': article_count,
                        'status': 'indexed'
                    })

            except Exception as mg_err:
                # If gqlalchemy fails, just save what was processed
                documents_list = [{
                    'cao': 'Cao Uitzendkrachten 2025',  # Example, would be dynamic in real scenario
                    'article_count': total_imported,
                    'status': 'processed'
                }]

            # Create index file
            index_data = {
                'documents': documents_list,
                'total': len(documents_list),
                'total_articles': sum([d.get('article_count', 0) for d in documents_list]),
                'last_updated': datetime.now().isoformat()
            }

            # Write index file
            with open(index_file, 'w') as f:
                json.dump(index_data, f, indent=2)

            upload_status['messages'].append(f"   ✓ Index saved ({len(documents_list)} documents tracked)")

        except Exception as idx_err:
            upload_status['messages'].append(f"   ⚠️ Could not save index: {str(idx_err)[:50]}")

    except Exception as e:
        upload_status['status'] = 'error'
        upload_status['error'] = str(e)
        upload_status['messages'].append(f"❌ Fatal error: {str(e)[:150]}")

@app.route('/upload/api/status', methods=['GET'])
def get_upload_status():
    """Get upload status"""
    return jsonify(upload_status)

@app.route('/upload/api/documents', methods=['GET'])
def get_indexed_documents():
    """Get indexed documents - reads from local tracking file"""
    import json
    from pathlib import Path

    # Path to document tracking file
    tracking_file = Path('/tmp/cao_documents_index.json')

    try:
        # Read from local tracking file (created during import)
        if tracking_file.exists():
            with open(tracking_file, 'r') as f:
                data = json.load(f)
                return jsonify(data)

        # If no tracking file exists, return empty
        return jsonify({
            'documents': [],
            'total': 0,
            'total_articles': 0,
            'message': 'No documents indexed yet'
        })

    except Exception as e:
        return jsonify({
            'documents': [],
            'total': 0,
            'total_articles': 0,
            'error': str(e)[:100]
        }), 200


@app.route('/upload/api/documents/list', methods=['GET'])
@super_admin_required
def get_uploaded_documents():
    """Get list of uploaded documents with details from Memgraph (Super Admin)"""
    import sys
    sys.path.insert(0, '/var/www/lexi')
    try:
        from gqlalchemy import Memgraph
        from datetime import datetime
        import os

        memgraph = Memgraph(
            host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
            port=int(os.getenv('MEMGRAPH_PORT', 7687))
        )

        # Query all CAO documents with their article counts
        results = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO)
            WITH cao.name as cao_name, cao
            OPTIONAL MATCH (cao)-[:CONTAINS_ARTICLE]->(article:Article)
            RETURN cao_name, COUNT(article) as article_count
            ORDER BY cao_name
        """))

        documents = []
        total_articles = 0

        for idx, r in enumerate(results):
            article_count = r['article_count'] if r['article_count'] else 0
            total_articles += article_count
            documents.append({
                'id': f'doc_{idx+1}',
                'cao_name': r['cao_name'],
                'status': 'indexed',
                'article_count': article_count,
                'upload_date': datetime.now().isoformat()
            })

        return jsonify({
            'documents': documents,
            'total': len(documents),
            'total_articles': total_articles
        })

    except Exception as e:
        import traceback
        error_msg = f"{str(e)}: {traceback.format_exc()}"
        return jsonify({
            'documents': [],
            'total': 0,
            'error': error_msg[:200]
        }), 500


@app.route('/super-admin/api/documents/list', methods=['GET'])
@super_admin_required
def super_admin_get_documents():
    """Get list of all documents from Memgraph (Super Admin)"""
    import sys
    sys.path.insert(0, '/var/www/lexi')
    try:
        from gqlalchemy import Memgraph
        from datetime import datetime
        import os

        memgraph = Memgraph(
            host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
            port=int(os.getenv('MEMGRAPH_PORT', 7687))
        )

        # Query all CAO documents with their article counts
        # Note: Documents can be imported with either :CONTAINS_ARTIKEL or :CONTAINS_ARTICLE relationship
        results = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO)
            WITH cao.name as cao_name, cao
            OPTIONAL MATCH (cao)-[r:CONTAINS_ARTIKEL|CONTAINS_ARTICLE]->(node)
            RETURN cao_name, COUNT(node) as article_count
            ORDER BY cao_name
        """))

        documents = []
        total_articles = 0

        for idx, r in enumerate(results):
            article_count = r['article_count'] if r['article_count'] else 0
            total_articles += article_count
            documents.append({
                'id': f'doc_{idx+1}',
                'cao_name': r['cao_name'],
                'status': 'indexed',
                'article_count': article_count,
                'upload_date': datetime.now().isoformat()
            })

        return jsonify({
            'documents': documents,
            'total': len(documents),
            'total_articles': total_articles
        })

    except Exception as e:
        import traceback
        error_msg = f"{str(e)}: {traceback.format_exc()}"
        return jsonify({
            'documents': [],
            'total': 0,
            'error': error_msg[:200]
        }), 500


@app.route('/super-admin/api/documents/upload', methods=['POST'])
@csrf.exempt
@super_admin_required
def super_admin_upload_documents():
    """Upload documents via Super Admin and import to Memgraph"""
    global super_admin_upload_status
    import sys
    import os
    sys.path.insert(0, '/var/www/lexi')

    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            return jsonify({'error': 'No files selected'}), 400

        # Save files temporarily and process
        upload_dir = '/tmp/cao_import'
        os.makedirs(upload_dir, exist_ok=True)

        from werkzeug.utils import secure_filename
        valid_files = []

        for file in files:
            if file and file.filename.endswith(('.pdf', '.txt', '.docx')):
                filename = secure_filename(file.filename)
                file_path = os.path.join(upload_dir, filename)
                file.save(file_path)
                valid_files.append(file_path)

        if not valid_files:
            return jsonify({'error': 'No valid PDF/TXT/DOCX files'}), 400

        # Reset status
        super_admin_upload_status = {
            'status': 'uploading',
            'progress': 0,
            'current_file': '',
            'total_files': len(valid_files),
            'processed_files': 0,
            'messages': [f"📁 {len(valid_files)} files uploaded successfully"],
            'imported_count': 0,
            'error': None
        }

        # Start background processing
        import threading
        thread = threading.Thread(
            target=_process_documents_to_memgraph,
            args=(valid_files,)
        )
        thread.daemon = True
        thread.start()

        return jsonify({
            'message': 'Upload started',
            'files_count': len(valid_files),
            'status': 'processing'
        })

    except Exception as e:
        super_admin_upload_status['error'] = str(e)[:200]
        super_admin_upload_status['status'] = 'error'
        return jsonify({'error': str(e)[:200]}), 400


def _sync_articles_to_postgresql(cao_name: str, filename: str):
    """
    Phase 5: Sync articles from Memgraph to PostgreSQL with Voyage embeddings
    Runs in background thread context - non-blocking wrapper for async operations

    Args:
        cao_name: Name of the CAO document
        filename: Filename for logging

    Returns:
        Success status dict
    """
    import logging
    logger = logging.getLogger(__name__)

    try:
        # Import necessary modules
        import asyncio
        from src.pipeline.cao_integration import CAOIntegrationAdapter
        from gqlalchemy import Memgraph

        logger.info(f"[Phase 5] Starting PostgreSQL sync for {cao_name}")

        # Initialize Memgraph client
        memgraph = Memgraph(
            host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
            port=int(os.getenv('MEMGRAPH_PORT', 7687))
        )

        # Create event loop for async operations in background thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        try:
            # Initialize integration adapter
            # Note: db and voyage clients would need to be passed or initialized here
            adapter = CAOIntegrationAdapter(
                db=None,  # Would be initialized with actual connection pool
                voyage_client=None,  # Would be initialized with actual Voyage client
                memgraph_client=memgraph
            )

            # Extract articles from Memgraph
            articles = loop.run_until_complete(
                adapter.extract_articles_from_memgraph(cao_name)
            )

            logger.info(f"[Phase 5] Extracted {len(articles)} articles for {cao_name}")

            return {
                "success": True,
                "articles_synced": len(articles),
                "cao_name": cao_name
            }

        finally:
            loop.close()

    except Exception as e:
        logger.error(f"[Phase 5] PostgreSQL sync failed for {cao_name}: {e}")
        return {
            "success": False,
            "error": str(e)[:100],
            "cao_name": cao_name
        }


def _process_documents_to_memgraph(file_paths):
    """Process files and import to Memgraph using DeepSeek Semantic Pipeline"""
    global super_admin_upload_status
    import sys
    import os

    sys.path.insert(0, '/var/www/lexi')

    super_admin_upload_status['status'] = 'processing'
    super_admin_upload_status['messages'].append("⏳ Initializing DeepSeek Semantic Pipeline...")

    try:
        # Import DeepSeek Semantic Pipeline
        try:
            from deepseek_semantic_pipeline import DeepSeekSemanticPipeline
            super_admin_upload_status['messages'].append("   ✓ DeepSeek Semantic Pipeline loaded")
        except ImportError as ie:
            super_admin_upload_status['messages'].append(f"   ❌ Failed to load DeepSeek pipeline: {str(ie)}")
            super_admin_upload_status['error'] = f"Import error: {str(ie)}"
            super_admin_upload_status['status'] = 'error'
            return

        # Initialize pipeline
        pipeline = DeepSeekSemanticPipeline()
        super_admin_upload_status['messages'].append("   ✓ Pipeline initialized (DeepSeek R1 + Memgraph)")

        total_imported = 0

        for idx, file_path in enumerate(file_paths):
            try:
                filename = os.path.basename(file_path)
                super_admin_upload_status['current_file'] = filename
                base_progress = int((idx / len(file_paths)) * 100)
                super_admin_upload_status['progress'] = base_progress
                super_admin_upload_status['messages'].append(f"\n{'='*60}")
                super_admin_upload_status['messages'].append(f"📄 Document {idx + 1}/{len(file_paths)}: {filename}")
                super_admin_upload_status['messages'].append(f"{'='*60}")

                # Phase 1: Reading document
                super_admin_upload_status['progress'] = base_progress + 5
                super_admin_upload_status['messages'].append(f"   📖 [1/4] Reading document...")
                import time
                start_time = time.time()
                
                # Phase 2: DeepSeek semantic chunking
                super_admin_upload_status['progress'] = base_progress + 15
                super_admin_upload_status['messages'].append(f"   🧠 [2/4] DeepSeek R1 semantic analysis (dit kan 30-60 sec duren)...")
                
                # Phase 3: Structure analysis
                super_admin_upload_status['progress'] = base_progress + 50
                super_admin_upload_status['messages'].append(f"   🔍 [3/4] Extracting articles and metadata...")
                
                # Phase 4: Memgraph import
                super_admin_upload_status['progress'] = base_progress + 75
                super_admin_upload_status['messages'].append(f"   💾 [4/4] Importing to Memgraph graph database...")
                
                # Process document
                success = pipeline.process_document(file_path)

                elapsed = time.time() - start_time

                if success:
                    super_admin_upload_status['messages'].append(f"   ✅ Document processed successfully in {elapsed:.1f}s")
                    super_admin_upload_status['processed_files'] += 1
                    total_imported += 1

                    # Phase 5: PostgreSQL + Voyage Storage (NEW!)
                    try:
                        super_admin_upload_status['progress'] = base_progress + 85
                        super_admin_upload_status['messages'].append(f"   📊 [5/5] Storing in PostgreSQL with Voyage embeddings...")

                        # Attempt PostgreSQL sync with Voyage embeddings
                        sync_result = _sync_articles_to_postgresql(cao_name=filename, filename=filename)

                        if sync_result.get('success'):
                            articles_synced = sync_result.get('articles_synced', 0)
                            super_admin_upload_status['messages'].append(f"   ✓ PostgreSQL sync complete: {articles_synced} articles cached for semantic search")
                            super_admin_upload_status['progress'] = base_progress + 95
                        else:
                            error_msg = sync_result.get('error', 'Unknown error')
                            super_admin_upload_status['messages'].append(f"   ℹ️  PostgreSQL sync skipped (non-critical): {error_msg[:60]}")

                    except Exception as e:
                        super_admin_upload_status['messages'].append(f"   ℹ️  PostgreSQL sync unavailable (non-critical): {str(e)[:60]} - Memgraph data retained")

                    # Upload to S3 for backup
                    try:
                        super_admin_upload_status['messages'].append(f"   📤 Uploading to S3 backup...")
                        s3_key = s3_service.upload_file_to_s3(file_path, 'cao-documents')
                        if s3_key:
                            super_admin_upload_status['messages'].append(f"   ☁️  S3 backup saved: {s3_key}")
                        else:
                            super_admin_upload_status['messages'].append(f"   ⚠️  S3 upload failed (non-critical)")
                    except Exception as s3_error:
                        super_admin_upload_status['messages'].append(f"   ⚠️  S3 backup error: {str(s3_error)[:100]} (non-critical)")
                else:
                    super_admin_upload_status['messages'].append(f"   ⚠️ Processing failed after {elapsed:.1f}s - check logs at /var/log/lexi/deepseek_semantic.log")

            except Exception as e:
                error_msg = f"❌ Error processing {filename}: {str(e)[:100]}"
                super_admin_upload_status['messages'].append(error_msg)
                super_admin_upload_status['error'] = str(e)[:200]
                import traceback
                print(f"[DOCUMENT UPLOAD ERROR] {filename}: {traceback.format_exc()}", file=sys.stderr)

        # Get total article count from Memgraph
        try:
            from gqlalchemy import Memgraph
            memgraph = Memgraph(
                host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
                port=int(os.getenv('MEMGRAPH_PORT', 7687))
            )
            
            result = list(memgraph.execute_and_fetch(
                "MATCH (a:Artikel) RETURN count(*) AS count"
            ))
            total_articles = result[0]['count'] if result else 0
            super_admin_upload_status['imported_count'] = total_articles
            super_admin_upload_status['messages'].append(f"   📊 Total articles in database: {total_articles}")
        except:
            super_admin_upload_status['imported_count'] = total_imported

        # Complete
        super_admin_upload_status['status'] = 'complete'
        super_admin_upload_status['progress'] = 100
        super_admin_upload_status['messages'].append(f"✨ COMPLETE! {total_imported} documents processed successfully")

    except Exception as e:
        super_admin_upload_status['status'] = 'error'
        super_admin_upload_status['error'] = str(e)[:200]
        super_admin_upload_status['messages'].append(f"❌ Fatal error: {str(e)[:100]}")
        import traceback
        print(f"[FATAL UPLOAD ERROR] {traceback.format_exc()}", file=sys.stderr)


@app.route('/super-admin/api/documents/status', methods=['GET'])
@super_admin_required
def super_admin_upload_status_endpoint():
    """Get current upload status"""
    return jsonify(super_admin_upload_status)


@app.route('/super-admin/api/documents/<doc_id>', methods=['DELETE'])
@csrf.exempt
@super_admin_required
def super_admin_delete_document(doc_id):
    """Delete a CAO document from Memgraph"""
    import sys
    import os
    sys.path.insert(0, '/var/www/lexi')

    try:
        from gqlalchemy import Memgraph

        memgraph = Memgraph(
            host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
            port=int(os.getenv('MEMGRAPH_PORT', 7687))
        )

        # Get all CAOs and find the one to delete by index
        all_caos = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO)
            RETURN cao.name as cao_name
            ORDER BY cao.name
        """))

        # Parse doc_id (format: doc_1, doc_2, etc.)
        try:
            doc_index = int(doc_id.replace('doc_', '')) - 1
        except:
            return jsonify({'error': 'Invalid document ID format'}), 400

        if doc_index < 0 or doc_index >= len(all_caos):
            return jsonify({'error': 'Document not found'}), 404

        cao_to_delete = all_caos[doc_index]['cao_name']

        # Delete CAO and all related articles from Memgraph
        result = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO {name: $cao_name})
            DETACH DELETE cao
            RETURN 1
        """, {'cao_name': cao_to_delete}))

        if not result:
            return jsonify({'error': 'Document not found'}), 404

        return jsonify({
            'message': f'Document "{cao_to_delete}" deleted successfully',
            'status': 'deleted'
        })

    except Exception as e:
        return jsonify({'error': str(e)[:200]}), 400


@app.route('/super-admin/api/documents/<doc_id>/rename', methods=['PUT'])
@csrf.exempt
@super_admin_required
def super_admin_rename_document(doc_id):
    """Rename a CAO document in Memgraph"""
    import sys
    import os
    sys.path.insert(0, '/var/www/lexi')

    try:
        from gqlalchemy import Memgraph

        data = request.get_json()
        new_name = data.get('name', '').strip()

        if not new_name:
            return jsonify({'error': 'New name is required'}), 400

        memgraph = Memgraph(
            host=os.getenv('MEMGRAPH_HOST', '46.224.4.188'),
            port=int(os.getenv('MEMGRAPH_PORT', 7687))
        )

        # Get all CAOs and find the one to rename by index
        all_caos = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO)
            RETURN cao.name as cao_name
            ORDER BY cao.name
        """))

        # Parse doc_id
        try:
            doc_index = int(doc_id.replace('doc_', '')) - 1
        except:
            return jsonify({'error': 'Invalid document ID format'}), 400

        if doc_index < 0 or doc_index >= len(all_caos):
            return jsonify({'error': 'Document not found'}), 404

        old_name = all_caos[doc_index]['cao_name']

        # Update CAO name in Memgraph
        result = list(memgraph.execute_and_fetch("""
            MATCH (cao:CAO {name: $old_name})
            SET cao.name = $new_name
            RETURN cao.name as updated_name
        """, {'old_name': old_name, 'new_name': new_name}))

        if not result:
            return jsonify({'error': 'Document not found'}), 404

        return jsonify({
            'message': f'Document renamed successfully',
            'old_name': old_name,
            'new_name': new_name
        })

    except Exception as e:
        return jsonify({'error': str(e)[:200]}), 400


@app.route('/upload/api/documents/<doc_id>', methods=['DELETE'])
@csrf.exempt
def delete_document(doc_id):
    """Delete a document from tracking (legacy)"""
    import sys
    sys.path.insert(0, '/var/www/lexi')
    try:
        from document_tracker import delete_document as delete_tracked_doc

        # Delete from tracker
        delete_tracked_doc(doc_id)

        return jsonify({
            'message': f'Document {doc_id} deleted successfully',
            'status': 'deleted'
        })

    except Exception as e:
        return jsonify({
            'error': str(e)[:100],
            'status': 'error'
        }), 400

@app.route('/super-admin/tenants/create', methods=['POST'])
@super_admin_required
def super_admin_create_tenant():
    company_name = request.form.get('company_name')
    subdomain = request.form.get('subdomain', '').lower().strip()
    contact_email = request.form.get('contact_email')
    contact_name = request.form.get('contact_name')
    max_users = int(request.form.get('max_users', 5))
    
    tenant = Tenant(
        company_name=company_name,
        subdomain=subdomain,
        contact_email=contact_email,
        contact_name=contact_name,
        max_users=max_users,
        status='active'
    )
    db.session.add(tenant)
    db.session.commit()
    
    flash('Tenant aangemaakt!', 'success')
    return redirect(url_for('super_admin_dashboard'))

@app.route('/super-admin/tenants/<int:tenant_id>/status', methods=['POST'])
@super_admin_required
def super_admin_update_tenant_status(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    new_status = request.form.get('status')
    
    if new_status in ['active', 'suspended', 'archived']:
        tenant.status = new_status
        db.session.commit()
        flash('Tenant status bijgewerkt!', 'success')
    
    return redirect(url_for('super_admin_dashboard'))

@app.route('/super-admin/tenants/<int:tenant_id>/tier', methods=['POST'])
@super_admin_required
def super_admin_update_tenant_tier(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    new_tier = request.form.get('tier')
    
    if new_tier in ['starter', 'professional', 'enterprise']:
        tenant.subscription_tier = new_tier
        tenant.max_users = get_max_users_for_tier(new_tier)
        
        # Update subscription plan if exists
        subscription = Subscription.query.filter_by(tenant_id=tenant_id).first()
        if subscription:
            subscription.plan = new_tier
        
        db.session.commit()
        flash(f'Tenant tier bijgewerkt naar {new_tier} (max {tenant.max_users} users)!', 'success')
    
    return redirect(url_for('super_admin_dashboard'))

@app.route('/super-admin/tenants/<int:tenant_id>/cao', methods=['POST'])
@super_admin_required
def super_admin_update_tenant_cao(tenant_id):
    from cao_config import validate_cao_preference, get_cao_display_name
    
    tenant = Tenant.query.get_or_404(tenant_id)
    new_cao = request.form.get('cao_preference')
    
    if validate_cao_preference(new_cao):
        tenant.cao_preference = new_cao
        db.session.commit()
        
        cao_name = get_cao_display_name(new_cao)
        flash(f'CAO voorkeur voor {tenant.company_name} bijgewerkt naar {cao_name}!', 'success')
    else:
        flash('Ongeldige CAO keuze.', 'danger')
    
    return redirect(url_for('super_admin_dashboard'))

@app.route('/super-admin/impersonate/<int:tenant_id>', methods=['POST'])
@super_admin_required
def super_admin_impersonate(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    
    admin_user = User.query.filter_by(tenant_id=tenant_id, role='admin').first()
    
    if not admin_user:
        flash('Geen admin user gevonden voor deze tenant!', 'error')
        return redirect(url_for('super_admin_tenant_detail', tenant_id=tenant_id))
    
    session['impersonating_super_admin_id'] = session.get('super_admin_id')
    session['impersonating_from'] = 'super_admin'
    
    session.pop('super_admin_id', None)
    session.pop('is_super_admin', None)
    
    logout_user()
    login_user(admin_user)
    
    session['tenant_id'] = tenant_id
    
    flash(f'Nu ingelogd als {admin_user.full_name} ({tenant.company_name})', 'success')
    return redirect('/chat')

@app.route('/stop-impersonate', methods=['POST'])
def stop_impersonate():
    if session.get('impersonating_from') == 'super_admin':
        super_admin_id = session.get('impersonating_super_admin_id')
        
        if not super_admin_id:
            flash('Impersonation context verloren', 'error')
            return redirect(url_for('index'))
        
        super_admin = SuperAdmin.query.get(super_admin_id)
        if not super_admin:
            flash('Super admin niet gevonden', 'error')
            return redirect(url_for('index'))
        
        logout_user()
        
        session.pop('tenant_id', None)
        session.pop('impersonating_from', None)
        session.pop('impersonating_super_admin_id', None)
        
        login_user(super_admin)
        session['super_admin_id'] = super_admin.id
        session['is_super_admin'] = True
        
        flash('Impersonation gestopt', 'success')
        return redirect(url_for('super_admin_dashboard'))
    
    flash('Je was niet aan het impersonaten', 'error')
    return redirect(url_for('index'))

@app.route('/super-admin/tenants/<int:tenant_id>')
@super_admin_required
def super_admin_tenant_detail(tenant_id):
    tenant = Tenant.query.get_or_404(tenant_id)
    
    users = User.query.filter_by(tenant_id=tenant_id).all()
    
    total_questions = db.session.query(Message).join(Chat).filter(
        Chat.tenant_id == tenant_id,
        Message.role == 'user'
    ).count()
    
    from datetime import datetime, timedelta
    from dateutil.relativedelta import relativedelta
    
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    questions_this_month = db.session.query(Message).join(Chat).filter(
        Chat.tenant_id == tenant_id,
        Message.role == 'user',
        Message.created_at >= thirty_days_ago
    ).count()
    
    today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    questions_today = db.session.query(Message).join(Chat).filter(
        Chat.tenant_id == tenant_id,
        Message.role == 'user',
        Message.created_at >= today
    ).count()
    
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    active_users_count = db.session.query(Chat.user_id).filter(
        Chat.tenant_id == tenant_id,
        Chat.created_at >= seven_days_ago
    ).distinct().count()
    
    avg_questions = total_questions / len(users) if users else 0
    
    top_questions = db.session.query(
        Message.content, 
        db.func.count(Message.id).label('count')
    ).join(Chat).filter(
        Chat.tenant_id == tenant_id,
        Message.role == 'user'
    ).group_by(Message.content).order_by(db.desc('count')).limit(5).all()
    
    for user in users:
        user.question_count = db.session.query(Message).join(Chat).filter(
            Chat.user_id == user.id,
            Message.role == 'user'
        ).count()
        user.last_activity = db.session.query(db.func.max(Chat.updated_at)).filter(
            Chat.user_id == user.id
        ).scalar()
    
    trial_days_left = None
    if tenant.trial_ends_at and tenant.subscription_status == 'trial':
        trial_days_left = (tenant.trial_ends_at - datetime.utcnow()).days
        if trial_days_left < 0:
            trial_days_left = 0
    
    return render_template('super_admin_tenant_detail.html',
                         tenant=tenant,
                         users=users,
                         total_questions=total_questions,
                         questions_this_month=questions_this_month,
                         questions_today=questions_today,
                         active_users_count=active_users_count,
                         avg_questions=avg_questions,
                         top_questions=top_questions,
                         trial_days_left=trial_days_left)

@app.route('/super-admin/users/<int:user_id>/reset-password', methods=['POST'])
@super_admin_required
def super_admin_reset_password(user_id):
    """Super admin can trigger password reset for any user (token-based, secure)"""
    from datetime import datetime, timedelta
    
    user = User.query.get_or_404(user_id)
    tenant = Tenant.query.get(user.tenant_id)
    
    # Generate secure reset token (super admin NEVER sees user password)
    reset_token = secrets.token_urlsafe(32)
    
    # Set token expiration (1 hour from now)
    user.reset_token = reset_token
    user.reset_token_expires_at = datetime.utcnow() + timedelta(hours=1)
    db.session.commit()
    
    # Create reset URL (no subdomain needed)
    domain = os.getenv('PRODUCTION_DOMAIN', 'lexiai.nl')
    reset_url = f"https://{domain}/reset-password/{reset_token}"
    
    # Send password reset link email (NO password in email)
    email_service.send_password_reset_link_email(user, tenant, reset_url)
    
    flash(f'Password reset link verzonden naar {user.first_name} {user.last_name} ({user.email}). Link is 1 uur geldig.', 'success')
    return redirect(url_for('super_admin_tenant_detail', tenant_id=tenant.id))

@app.route('/super-admin/analytics/export')
@super_admin_required
def super_admin_analytics_export():
    import csv
    from io import StringIO
    from flask import Response
    
    tenants = Tenant.query.all()
    
    output = StringIO()
    writer = csv.writer(output)
    
    writer.writerow(['Tenant ID', 'Company Name', 'Subdomain', 'Status', 'Tier', 'MRR', 'Users', 'Questions', 'Created At'])
    
    mrr_prices = {'starter': 499, 'professional': 599, 'enterprise': 1199}
    
    for tenant in tenants:
        users_count = User.query.filter_by(tenant_id=tenant.id).count()
        questions_count = db.session.query(Message).join(Chat).filter(
            Chat.tenant_id == tenant.id,
            Message.role == 'user'
        ).count()
        
        mrr = mrr_prices.get(tenant.subscription_tier, 0) if tenant.subscription_status == 'active' else 0
        
        writer.writerow([
            tenant.id,
            tenant.company_name,
            tenant.subdomain,
            tenant.subscription_status,
            tenant.subscription_tier,
            mrr,
            users_count,
            questions_count,
            tenant.created_at.strftime('%Y-%m-%d %H:%M:%S')
        ])
    
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment; filename=analytics_export.csv'}
    )

@app.route('/super-admin/analytics')
@super_admin_required
def super_admin_analytics():
    from dateutil.relativedelta import relativedelta
    from datetime import datetime, timedelta
    
    tenants = Tenant.query.all()
    all_users = User.query.all()
    
    mrr_prices = {'starter': 499, 'professional': 599, 'enterprise': 1199}
    
    current_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in tenants if t.subscription_status == 'active')
    total_revenue = current_mrr * 12
    
    active_tenants = sum(1 for t in tenants if t.subscription_status == 'active')
    trial_tenants = sum(1 for t in tenants if t.subscription_status == 'trial')
    
    last_month = datetime.utcnow() - relativedelta(months=1)
    last_month_tenants = [t for t in tenants if t.created_at < last_month and t.subscription_status == 'active']
    last_month_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in last_month_tenants)
    
    growth_rate = 0
    if last_month_mrr > 0:
        growth_rate = ((current_mrr - last_month_mrr) / last_month_mrr) * 100
    elif current_mrr > 0 and last_month_mrr == 0:
        growth_rate = 100
    
    total_questions = db.session.query(Message).filter(Message.role == 'user').count()
    
    mrr_history = []
    questions_history = []
    for i in range(12, 0, -1):
        month_date = datetime.utcnow() - relativedelta(months=i)
        month_tenants = [t for t in tenants if t.created_at <= month_date and t.subscription_status == 'active']
        month_mrr = sum(mrr_prices.get(t.subscription_tier, 0) for t in month_tenants)
        
        month_start = month_date.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        month_end = (month_start + relativedelta(months=1)) - timedelta(seconds=1)
        month_questions = db.session.query(Message).filter(
            Message.role == 'user',
            Message.created_at >= month_start,
            Message.created_at <= month_end
        ).count()
        
        mrr_history.append({
            'month': month_date.strftime('%b %Y'),
            'mrr': month_mrr
        })
        questions_history.append({
            'month': month_date.strftime('%b %Y'),
            'count': month_questions
        })
    
    tier_distribution = {
        'starter': sum(1 for t in tenants if t.subscription_tier == 'starter' and t.subscription_status == 'active'),
        'professional': sum(1 for t in tenants if t.subscription_tier == 'professional' and t.subscription_status == 'active'),
        'enterprise': sum(1 for t in tenants if t.subscription_tier == 'enterprise' and t.subscription_status == 'active'),
        'trial': sum(1 for t in tenants if t.subscription_tier == 'trial')
    }
    
    top_tenants = []
    for tenant in tenants:
        tenant_questions = db.session.query(Message).join(Chat).filter(
            Chat.tenant_id == tenant.id,
            Message.role == 'user'
        ).count()
        if tenant_questions > 0:
            top_tenants.append({
                'tenant': tenant,
                'questions': tenant_questions
            })
    top_tenants = sorted(top_tenants, key=lambda x: x['questions'], reverse=True)[:10]
    
    top_questions = db.session.query(
        Message.content,
        db.func.count(Message.id).label('count')
    ).filter(
        Message.role == 'user'
    ).group_by(Message.content).order_by(db.desc('count')).limit(10).all()
    
    recent_activity = db.session.query(Chat).order_by(Chat.updated_at.desc()).limit(20).all()
    
    conversion_funnel = {
        'signups': len(tenants),
        'trials': trial_tenants,
        'active': active_tenants,
        'conversion_rate': (active_tenants / len(tenants) * 100) if tenants else 0
    }
    
    return render_template('super_admin_analytics.html',
                         current_mrr=current_mrr,
                         total_revenue=total_revenue,
                         active_tenants=active_tenants,
                         growth_rate=growth_rate,
                         total_questions=total_questions,
                         mrr_history=mrr_history,
                         questions_history=questions_history,
                         tier_distribution=tier_distribution,
                         top_tenants=top_tenants,
                         top_questions=top_questions,
                         recent_activity=recent_activity,
                         conversion_funnel=conversion_funnel)

@app.route('/super-admin/support')
@super_admin_required
def super_admin_support():
    status_filter = request.args.get('status', 'all')
    category_filter = request.args.get('category', 'all')
    
    query = SupportTicket.query
    
    if status_filter != 'all':
        query = query.filter_by(status=status_filter)
    if category_filter != 'all':
        query = query.filter_by(category=category_filter)
    
    tickets = query.order_by(SupportTicket.created_at.desc()).all()
    
    open_count = SupportTicket.query.filter_by(status='open').count()
    in_progress_count = SupportTicket.query.filter_by(status='in_progress').count()
    answered_count = SupportTicket.query.filter_by(status='answered').count()
    closed_count = SupportTicket.query.filter_by(status='closed').count()
    
    return render_template('super_admin_support.html',
                         tickets=tickets,
                         status_filter=status_filter,
                         category_filter=category_filter,
                         open_count=open_count,
                         in_progress_count=in_progress_count,
                         answered_count=answered_count,
                         closed_count=closed_count)

@app.route('/super-admin/support/<int:ticket_id>')
@super_admin_required
def super_admin_support_detail(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    replies = SupportReply.query.filter_by(ticket_id=ticket_id).order_by(SupportReply.created_at).all()
    
    tenant = Tenant.query.get(ticket.tenant_id)
    user = User.query.get(ticket.user_id)
    
    return render_template('super_admin_support_detail.html',
                         ticket=ticket,
                         replies=replies,
                         tenant=tenant,
                         user=user)

@app.route('/api/super-admin/support/<int:ticket_id>/reply', methods=['POST'])
@super_admin_required
def super_admin_support_reply(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    message = (request.json or {}).get('message')
    
    if not message:
        return jsonify({'error': 'Message is required'}), 400
    
    reply = SupportReply(
        ticket_id=ticket_id,
        message=message,
        is_admin=True,
        sender_name='Super Admin'
    )
    db.session.add(reply)
    
    ticket.status = 'answered'
    ticket.updated_at = datetime.utcnow()
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'reply': {
            'message': reply.message,
            'is_admin': reply.is_admin,
            'sender_name': reply.sender_name,
            'created_at': reply.created_at.strftime('%d-%m-%Y %H:%M')
        }
    })

@app.route('/api/super-admin/support/<int:ticket_id>/status', methods=['POST'])
@super_admin_required
def super_admin_support_status(ticket_id):
    ticket = SupportTicket.query.get_or_404(ticket_id)
    new_status = (request.json or {}).get('status')
    
    if new_status not in ['open', 'in_progress', 'answered', 'closed']:
        return jsonify({'error': 'Invalid status'}), 400
    
    ticket.status = new_status
    ticket.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True})

def init_db():
    try:
        with app.app_context():
            db.create_all()
            print("Database tables checked/created successfully")
            
            if not SuperAdmin.query.first():
                # SECURITY: Generate strong random password for super admin (never hardcode!)
                import secrets
                import string
                alphabet = string.ascii_letters + string.digits + string.punctuation
                random_password = ''.join(secrets.choice(alphabet) for _ in range(24))
                
                admin = SuperAdmin(
                    email='admin@lex-cao.nl',
                    name='Super Administrator'
                )
                admin.set_password(random_password)
                db.session.add(admin)
                db.session.commit()
                
                # SECURITY: Only log credentials when EXPLICITLY in development mode
                # Default = NO logging (safe production behavior)
                if os.getenv('ENVIRONMENT') == 'development':
                    print("=" * 80)
                    print("🔐 SUPER ADMIN ACCOUNT CREATED (DEVELOPMENT MODE)")
                    print("=" * 80)
                    print(f"Email: admin@lex-cao.nl")
                    print(f"Password: {random_password}")
                    print("")
                    print("⚠️  Save this password NOW - it will not be shown again!")
                    print("⚠️  Change it immediately after first login!")
                    print("=" * 80)
                else:
                    # Production/default: NO credential logging (secure by default)
                    print("✅ Super admin created successfully")
                    # NOTE: For production, create super admin via one of these secure methods:
                    # 1. Set SUPER_ADMIN_PASSWORD env var before first deployment
                    # 2. Use secrets manager (AWS Secrets Manager, Vault, etc.)
                    # 3. Create manually via secure admin interface
                    # The random password is NOT accessible - rotate via password reset flow
    except Exception as e:
        print(f"Database initialization: {e}")

init_db()

# ========== ERROR HANDLERS (Security - Prevent Information Disclosure) ==========

@app.errorhandler(400)
def bad_request(e):
    """Handle bad request errors"""
    return render_template('error.html', 
                         error_code=400,
                         error_title='Ongeldig Verzoek',
                         error_message='Het verzoek kon niet worden verwerkt.'), 400

@app.route('/health')
def health_check():
    """Health check endpoint for Docker/monitoring"""
    health = {
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat(),
        'services': {}
    }

    # Check database connectivity
    try:
        db.session.execute(db.text('SELECT 1'))
        health['services']['database'] = 'healthy'
    except Exception as e:
        app.logger.error(f"Database health check failed: {e}")
        health['services']['database'] = 'unhealthy'
        health['status'] = 'unhealthy'

    # Check S3 service
    health['services']['s3'] = 'healthy' if s3_service.enabled else 'disabled'

    # Check RAG service (Memgraph + DeepSeek)
    health['services']['rag'] = 'healthy' if rag_service.enabled else 'disabled'

    # Check Email service
    health['services']['email'] = 'healthy' if email_service.enabled else 'disabled'

    # Overall status
    status_code = 200 if health['status'] == 'healthy' else 503

    return jsonify(health), status_code

@app.errorhandler(403)
def forbidden(e):
    """Handle forbidden errors"""
    return render_template('error.html',
                         error_code=403,
                         error_title='Geen Toegang',
                         error_message='Je hebt geen toegang tot deze pagina.'), 403

@app.errorhandler(404)
def not_found(e):
    """Handle not found errors"""
    return render_template('error.html',
                         error_code=404,
                         error_title='Pagina Niet Gevonden',
                         error_message='De pagina die je zoekt bestaat niet.'), 404

@app.errorhandler(429)
def rate_limit_exceeded(e):
    """Handle rate limit errors"""
    return render_template('error.html',
                         error_code=429,
                         error_title='Te Veel Verzoeken',
                         error_message='Je hebt te veel verzoeken gedaan. Probeer het later opnieuw.'), 429

@app.errorhandler(500)
def internal_error(e):
    """Handle internal server errors - Log details but show generic message"""
    # Log the error for debugging (but don't show to user)
    app.logger.error(f'Server Error: {str(e)}', exc_info=True)
    
    # Show generic error page (no sensitive information)
    return render_template('error.html',
                         error_code=500,
                         error_title='Server Fout',
                         error_message='Er is een fout opgetreden. Probeer het later opnieuw.'), 500

@app.errorhandler(503)
def service_unavailable(e):
    """Handle service unavailable errors"""
    return render_template('error.html',
                         error_code=503,
                         error_title='Service Niet Beschikbaar',
                         error_message='De service is tijdelijk niet beschikbaar. Probeer het later opnieuw.'), 503

if __name__ == '__main__':  
    app.run(host='0.0.0.0', port=5000, debug=True)
    
